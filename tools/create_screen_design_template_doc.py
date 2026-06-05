from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from create_split_design_docs import ASSET_DIR, SCREENS


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "MOMent_화면설계서_양식반영_수정본.docx"

FONT = "맑은 고딕"
AUTHOR = "정용철"

INK = "2D1B33"
MUTED = "7A6673"
PINK = "C94E70"
PINK_DARK = "A83D5B"
PINK_SOFT = "FFF1F5"
PINK_LIGHT = "FBE7EE"
PURPLE_SOFT = "F5F0FF"
GRID = "E7D7DF"
WHITE = "FFFFFF"


def set_run(run, size=9, bold=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_cell_text(cell, text, size=9, bold=False, color=INK, align=None, after=2):
    p = clear(cell) if not cell.text else cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_sidebar_item(cell, title, value, compact=False):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(5 if compact else 7)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    set_run(run, size=10.5, bold=True, color=PINK_DARK)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4 if compact else 6)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(value)
    set_run(run, size=7.6 if compact else 8.0, color=INK)


def normalize_uc(uc):
    return str(uc).replace("UC_", "UC-")


def feature_action(screen_id, no, name):
    actions = {
        "screen_home_01": {
            "1": "사용자가 서비스명과 앱 목적을 확인한다.",
            "2": "사용자가 핵심 기능 태그를 보고 서비스 범위를 파악한다.",
            "3": "사용자가 로그인 버튼을 누른다.",
            "4": "사용자가 회원가입 버튼을 누른다.",
            "5": "사용자가 정보 출처 안내를 확인한다.",
        },
        "screen_register_02": {
            "1": "사용자가 임산부 또는 보호자 역할을 선택한다.",
            "2": "사용자가 이메일과 비밀번호를 입력한다.",
            "3": "사용자가 이름과 태명을 입력한다.",
            "4": "사용자가 임신 시작 날짜를 캘린더에서 선택한다.",
            "5": "보호자 사용자가 임산부 연결 코드를 입력한다.",
            "6": "사용자가 가입 완료 버튼을 누른다.",
        },
        "screen_dashboard_03": {
            "1": "사용자가 현재 주차와 태명 요약을 확인한다.",
            "2": "사용자가 자동 전환되는 오늘의 팁/혜택 배너를 확인한다.",
            "3": "사용자가 오늘의 상태 체크 카드로 이동한다.",
            "4": "보호자가 오늘 생성된 미션 카드로 이동한다.",
            "5": "사용자가 원하는 기능 카드를 선택한다.",
            "6": "사용자가 하단 탭으로 주요 화면을 전환한다.",
        },
        "screen_status_check_04": {
            "1": "임산부가 오늘 느낀 증상을 복수 선택한다.",
            "2": "임산부가 현재 감정을 선택한다.",
            "3": "임산부가 상태 저장 버튼을 누른다.",
            "4": "사용자가 생성된 미션 요약을 확인한다.",
            "5": "보호자가 설정한 케어 선호가 자동 적용된다.",
        },
        "screen_diary_05": {
            "1": "사용자가 일기 작성 버튼을 눌러 입력 화면을 연다.",
            "2": "사용자가 직접 감정을 선택한다.",
            "3": "사용자가 AI 자동 감정 분석 버튼을 누른다.",
            "4": "사용자가 저장된 일기 목록을 날짜순으로 확인한다.",
            "5": "사용자가 양쪽 답변이 완료된 스몰토크 기록을 확인한다.",
            "6": "사용자가 카드를 눌러 상세 내용을 펼치거나 접는다.",
        },
        "screen_ai_recommend_06": {
            "1": "사용자가 현재 임신 주차와 태아 발달 요약을 확인한다.",
            "2": "사용자가 추천 식품 목록을 확인한다.",
            "3": "사용자가 권장 활동 목록을 확인한다.",
            "4": "사용자가 주의사항과 위험 신호를 확인한다.",
            "5": "사용자가 최근 다이어리 반영 여부를 확인한다.",
            "6": "사용자가 콘텐츠 탭을 전환한다.",
        },
        "screen_info_benefits_07": {
            "1": "사용자가 정보 카테고리 탭을 선택한다.",
            "2": "사용자가 검증 정보 카드를 확인한다.",
            "3": "사용자가 받을 수 있는 정부 혜택 카드를 확인한다.",
            "4": "사용자가 출처 링크를 눌러 원문 페이지로 이동한다.",
            "5": "사용자가 의학정보 챗봇 상담 영역을 연다.",
        },
        "screen_community_08": {
            "1": "사용자가 커뮤니티 게시글 목록을 확인한다.",
            "2": "사용자가 작성자 역할과 임신 시기 정보를 확인한다.",
            "3": "사용자가 새 게시글을 작성한다.",
            "4": "사용자가 댓글을 열람하거나 작성한다.",
            "5": "사용자가 좋아요 버튼을 눌러 토글한다.",
            "6": "사용자가 내 글 또는 내 댓글 화면으로 이동한다.",
        },
        "screen_appliance_09": {
            "1": "사용자가 AI 추천 가전 설정을 확인한다.",
            "2": "사용자가 가전별 현재 설정 카드를 확인한다.",
            "3": "사용자가 변경된 설정을 저장한다.",
            "4": "사용자가 Arduino 연결 상태를 확인한다.",
            "5": "사용자가 시연 동기화 버튼을 누른다.",
            "6": "사용자가 LG 가전 링크로 이동한다.",
        },
        "screen_smalltalk_10": {
            "1": "사용자가 오늘의 질문을 확인한다.",
            "2": "사용자가 자신의 답변을 입력한다.",
            "3": "사용자가 상대방 답변 완료 여부를 확인한다.",
            "4": "사용자가 답변 저장 버튼을 누른다.",
            "5": "두 사용자 답변이 모두 완료되면 다이어리 표시 조건이 충족된다.",
        },
        "screen_profile_settings_11": {
            "1": "사용자가 프로필 요약 정보를 확인한다.",
            "2": "사용자가 연결 코드를 확인하거나 공유한다.",
            "3": "보호자가 케어 선호 설정을 변경한다.",
            "4": "사용자가 로그아웃 버튼을 누른다.",
            "5": "사용자 역할에 따라 표시되는 메뉴를 확인한다.",
        },
        "screen_admin_12": {
            "1": "관리자가 전체 회원과 오늘 접속자 통계를 확인한다.",
            "2": "관리자가 회원 목록을 확인한다.",
            "3": "관리자가 특정 회원으로 대리접속한다.",
            "4": "관리자가 게시글, 댓글, 좋아요 수를 확인하고 삭제한다.",
            "5": "관리자가 커뮤니티 분석 버튼을 눌러 키워드를 분석한다.",
            "6": "관리자가 불용어를 입력하고 재분석한다.",
            "7": "관리자가 사용자 평균 가전 세팅 값을 확인한다.",
        },
    }
    return actions.get(screen_id, {}).get(no, f"사용자가 {name} 요소를 조작하거나 확인한다.")


def add_screen_page(doc, screen, index):
    if index > 0:
        doc.add_section(WD_SECTION.NEW_PAGE)

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    shell = doc.add_table(rows=1, cols=2)
    shell.alignment = WD_TABLE_ALIGNMENT.CENTER
    shell.autofit = False
    set_borders(shell, color=WHITE, size="2")

    left, right = shell.rows[0].cells
    set_width(left, 2600)
    set_width(right, 10550)
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade(left, PINK_SOFT)

    clear(left)
    add_sidebar_item(left, "화면 ID", screen["id"])
    add_sidebar_item(left, "화면 이름", screen["name"])
    add_sidebar_item(left, "관련 유스케이스 ID", normalize_uc(screen["uc"]))
    add_sidebar_item(left, "화면 개요", screen["overview"], compact=True)
    add_sidebar_item(left, "메뉴 경로", screen["path"])
    add_sidebar_item(left, "작성자", AUTHOR)

    clear(right)
    title = right.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run(f"{screen['name']} 화면 설계")
    set_run(run, size=18, bold=True, color=PINK)

    screenshot_block = right.add_table(rows=2, cols=1)
    screenshot_block.alignment = WD_TABLE_ALIGNMENT.CENTER
    screenshot_block.autofit = False
    set_borders(screenshot_block, color=GRID, size="5")
    screen_title = screenshot_block.rows[0].cells[0]
    screen_body = screenshot_block.rows[1].cells[0]
    shade(screen_title, PURPLE_SOFT)
    shade(screen_body, WHITE)
    add_cell_text(screen_title, "화면 설계", size=12, bold=True, color=PINK_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    p = clear(screen_body)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = ASSET_DIR / screen["image"]
    if img_path.exists():
        p.add_run().add_picture(str(img_path), width=Cm(6.4))
    else:
        run = p.add_run(f"화면 이미지 없음: {screen['image']}")
        set_run(run, size=9, bold=True, color=PINK)

    desc = right.add_table(rows=2, cols=1)
    desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    desc.autofit = False
    set_borders(desc, color=GRID, size="5")
    header = desc.rows[0].cells[0]
    body = desc.rows[1].cells[0]
    shade(header, PINK)
    shade(body, "FFF9FB")
    add_cell_text(header, "화면 설명(기능)", size=12.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    clear(body)
    f_table = body.add_table(rows=1, cols=4)
    f_table.autofit = False
    set_borders(f_table, color=GRID, size="4")
    headers = ["번호", "기능명", "사용자 액션(Input)", "처리/출력(Output)"]
    widths = [650, 1900, 3950, 4050]
    for i, text in enumerate(headers):
        cell = f_table.rows[0].cells[i]
        set_width(cell, widths[i])
        shade(cell, PINK_LIGHT)
        add_cell_text(cell, text, size=7.5, bold=True, color=PINK_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    for no, name, output in screen["features"]:
        cells = f_table.add_row().cells
        values = [no, name, feature_action(screen["id"], no, name), output]
        for i, value in enumerate(values):
            set_width(cells[i], widths[i])
            add_cell_text(cells[i], value, size=6.9, color=INK, after=0)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("유의사항: ")
    set_run(run, size=7.2, bold=True, color=PINK_DARK)
    run = p.add_run(screen["notes"])
    set_run(run, size=7.0, color=MUTED)


def build():
    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for index, screen in enumerate(SCREENS):
        add_screen_page(doc, screen, index)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
