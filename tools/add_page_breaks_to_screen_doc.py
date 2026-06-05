from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "MOMent_화면설계서_프로젝트뷰기준_최종본_그림삭제_압축수정.docx"
OUTPUT = ROOT / "MOMent_화면설계서_프로젝트뷰기준_최종본_그림삭제_페이지분리수정.docx"


def clear_paragraph_runs(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def compact_doc(doc):
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            if run.font.size is None or run.font.size.pt > 7.2:
                run.font.size = Pt(7.2)

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = "맑은 고딕"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
                        if run.font.size is None or run.font.size.pt > 6.8:
                            run.font.size = Pt(6.8)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document(SOURCE)
    compact_doc(doc)

    body_children = list(doc._body._element)
    # Document pattern is screen table followed by an empty paragraph.
    # Put a hard page break in each following paragraph except after the last screen table.
    table_indices = [idx for idx, el in enumerate(body_children) if el.tag == qn("w:tbl")]
    for tbl_index in table_indices[:-1]:
        if tbl_index + 1 < len(body_children) and body_children[tbl_index + 1].tag == qn("w:p"):
            paragraph = None
            target_el = body_children[tbl_index + 1]
            for p in doc.paragraphs:
                if p._p is target_el:
                    paragraph = p
                    break
            if paragraph is not None:
                clear_paragraph_runs(paragraph)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.add_run().add_break(WD_BREAK.PAGE)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
