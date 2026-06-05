from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "doc_assets" / "screen_design_full"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = ROOT / "MOMent_화면설계서_프로젝트뷰기준_최종본.docx"

FONT = "맑은 고딕"
AUTHOR = "팀원/정용철"

INK = "2D1B33"
MUTED = "7A6673"
PINK = "C94E70"
PINK_DARK = "A83D5B"
PINK_SOFT = "FFF1F5"
PINK_LIGHT = "FBE7EE"
PURPLE_SOFT = "F7F2FF"
GRID = "E7D7DF"
WHITE = "FFFFFF"


def pil_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/malgunbd.ttf") if bold else Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/NanumGothicBold.ttf") if bold else Path("C:/Windows/Fonts/NanumGothic.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


PF_TITLE = pil_font(28, True)
PF_H = pil_font(20, True)
PF_BODY = pil_font(16)
PF_SMALL = pil_font(13)


def wrap_text(draw, text, font, max_width):
    lines = []
    current = ""
    for token in str(text).split():
        trial = token if not current else f"{current} {token}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = token
    if current:
        lines.append(current)
    return lines


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_mock_screen(spec):
    path = ASSET_DIR / spec["image"]
    img = Image.new("RGB", (430, 860), "#FFF7FA")
    d = ImageDraw.Draw(img)
    d.text((26, 22), "9:41", font=PF_SMALL, fill=f"#{INK}")
    d.text((330, 22), "LTE 100%", font=PF_SMALL, fill=f"#{INK}")

    rounded(d, (20, 62, 410, 126), 24, "#FFFFFF", "#F0DDE5", 2)
    d.text((40, 82), spec["name"], font=PF_TITLE, fill=f"#{INK}")
    d.text((40, 112), spec["screen_file"], font=PF_SMALL, fill=f"#{PINK}")

    y = 150
    for idx, feature in enumerate(spec["features"][:6], start=1):
        h = 82
        rounded(d, (24, y, 406, y + h), 18, "#FFFFFF", "#EADCE3", 2)
        rounded(d, (42, y + 18, 74, y + 50), 16, "#FCE7F0")
        d.text((53, y + 24), str(idx), font=PF_SMALL, fill=f"#{PINK_DARK}")
        d.text((88, y + 15), feature["name"], font=PF_H, fill=f"#{INK}")
        for line_no, line in enumerate(wrap_text(d, feature["output"], PF_SMALL, 285)[:2]):
            d.text((88, y + 43 + line_no * 17), line, font=PF_SMALL, fill=f"#{MUTED}")
        y += h + 14
        if y > 690:
            break

    rounded(d, (24, 724, 406, 770), 20, "#C94E70")
    d.text((54, 738), spec["route"], font=PF_SMALL, fill="#FFFFFF")
    d.rectangle((0, 792, 430, 860), fill="#FFFFFF")
    nav = ["홈", "가전", "일기", "내정보", "설정"]
    for i, item in enumerate(nav):
        x = 43 + i * 86
        color = "#C94E70" if item in spec.get("active_nav", "") else "#9B8C95"
        d.ellipse((x - 8, 808, x + 8, 824), fill=color)
        d.text((x - 18, 832), item, font=PF_SMALL, fill=color)
    img.save(path)
    return path


SCREENS = [
    {
        "id": "screen_home_01",
        "name": "MOMent 시작 화면",
        "screen_file": "App.tsx / HomeView",
        "uc": "UC-01, UC-02",
        "overview": "서비스 최초 진입 화면으로 로그인과 회원가입을 선택한다.",
        "route": "앱 실행 > 시작 화면",
        "active_nav": "홈",
        "image": "01_home.png",
        "features": [
            {"name": "브랜드 영역", "action": "사용자가 서비스명과 앱 목적을 확인한다.", "output": "MOMent 서비스명과 임신 케어 핵심 메시지를 표시한다."},
            {"name": "기능 태그", "action": "사용자가 주요 기능 범위를 확인한다.", "output": "가전, 보호자 케어, AI 추천, 신뢰 정보, 커뮤니티 기능을 요약한다."},
            {"name": "로그인 버튼", "action": "사용자가 로그인 버튼을 누른다.", "output": "LoginView로 이동한다."},
            {"name": "회원가입 버튼", "action": "사용자가 회원가입 버튼을 누른다.", "output": "RegisterView로 이동한다."},
        ],
        "notes": "비로그인 사용자는 홈, 로그인, 회원가입 화면만 접근 가능하다.",
    },
    {
        "id": "screen_login_02",
        "name": "로그인 화면",
        "screen_file": "LoginView.tsx",
        "uc": "UC-02",
        "overview": "이메일과 비밀번호로 사용자를 인증하고 역할별 초기 화면으로 분기한다.",
        "route": "시작 화면 > 로그인",
        "active_nav": "홈",
        "image": "02_login.png",
        "features": [
            {"name": "이메일 입력", "action": "사용자가 가입 이메일을 입력한다.", "output": "로그인 API 요청의 email 값으로 전달한다."},
            {"name": "비밀번호 입력", "action": "사용자가 비밀번호를 입력한다.", "output": "로그인 API 요청의 password 값으로 전달한다."},
            {"name": "로그인 처리", "action": "사용자가 로그인 버튼을 누른다.", "output": "성공 시 user_id, role, 임신 정보, 연결 정보를 저장한다."},
            {"name": "역할 분기", "action": "서버가 사용자 role을 반환한다.", "output": "ADMIN은 관리자 화면, 일반 사용자는 대시보드로 이동한다."},
            {"name": "회원가입 이동", "action": "사용자가 회원가입 링크를 누른다.", "output": "RegisterView로 이동한다."},
        ],
        "notes": "서버 응답이 늦을 경우 로그인 요청에 타임아웃을 적용하고 중복 클릭을 막는다.",
    },
    {
        "id": "screen_register_03",
        "name": "회원가입 화면",
        "screen_file": "RegisterView.tsx",
        "uc": "UC-01",
        "overview": "임산부/보호자 역할에 따라 계정 정보와 연결 정보를 입력한다.",
        "route": "시작 화면 > 회원가입",
        "active_nav": "홈",
        "image": "03_register.png",
        "features": [
            {"name": "역할 선택", "action": "사용자가 임산부 또는 보호자를 선택한다.", "output": "역할에 따라 임신 시작일, 태명, 연결 코드 입력 항목을 분기한다."},
            {"name": "계정 입력", "action": "사용자가 이메일, 비밀번호, 이름을 입력한다.", "output": "USERS 생성 요청의 기본 데이터로 사용한다."},
            {"name": "임신 시작일", "action": "임산부 사용자가 날짜를 선택한다.", "output": "오늘 기준 280일 전부터 오늘까지 허용하고 주차 계산 기준으로 저장한다."},
            {"name": "연결 코드", "action": "보호자가 임산부 연결 코드를 입력한다.", "output": "parent_user_id와 connection_code를 통해 가족 연결을 구성한다."},
            {"name": "가입 완료", "action": "사용자가 가입 버튼을 누른다.", "output": "성공 시 로그인 상태로 전환한다."},
        ],
        "notes": "임신 시작일은 미래 날짜가 입력되지 않도록 제한한다.",
    },
    {
        "id": "screen_dashboard_04",
        "name": "대시보드",
        "screen_file": "DashboardView.tsx",
        "uc": "UC-03, UC-04, UC-05, UC-09, UC-16, UC-19",
        "overview": "주차, 오늘의 팁, 혜택, 주요 기능 진입점을 제공하는 메인 화면이다.",
        "route": "로그인 > 대시보드",
        "active_nav": "홈",
        "image": "04_dashboard.png",
        "features": [
            {"name": "사용자 주차 요약", "action": "사용자가 태명과 현재 주차를 확인한다.", "output": "USERS.pregnancy_start_date 기준 주차를 계산해 표시한다."},
            {"name": "오늘의 팁", "action": "사용자가 자동 전환 배너를 확인한다.", "output": "주차별 팁과 정부 혜택 추천을 슬라이드 형태로 보여준다."},
            {"name": "상태 체크 진입", "action": "임산부가 상태 체크 카드를 누른다.", "output": "DiscomfortView로 이동한다."},
            {"name": "보호자 미션 진입", "action": "보호자가 미션 카드를 누른다.", "output": "MissionView로 이동한다."},
            {"name": "주요 기능 카드", "action": "사용자가 원하는 기능을 선택한다.", "output": "AI 추천, 신뢰 정보, 커뮤니티, 스몰토크 등으로 이동한다."},
        ],
        "notes": "보호자 계정은 연결된 임산부 기준 데이터를 조회한다.",
    },
    {
        "id": "screen_status_check_05",
        "name": "오늘의 상태 체크",
        "screen_file": "DiscomfortView.tsx",
        "uc": "UC-04, UC-05, UC-06",
        "overview": "임산부가 오늘의 증상과 감정을 저장하면 보호자 미션이 생성된다.",
        "route": "대시보드 > 오늘의 상태 체크",
        "active_nav": "홈",
        "image": "05_status_check.png",
        "features": [
            {"name": "증상 선택", "action": "임산부가 오늘의 증상을 복수 선택한다.", "output": "PREGNANCY_STATUS_CHECKS.symptoms에 JSON 형태로 저장한다."},
            {"name": "감정 선택", "action": "임산부가 현재 감정을 선택한다.", "output": "PREGNANCY_STATUS_CHECKS.emotions에 JSON 형태로 저장한다."},
            {"name": "상태 저장", "action": "임산부가 저장 버튼을 누른다.", "output": "상태 체크 생성 후 보호자 미션 생성 로직을 호출한다."},
            {"name": "미션 생성", "action": "서버가 연결된 보호자를 조회한다.", "output": "GUARDIAN_MISSIONS에 상태 기반 미션을 생성한다."},
        ],
        "notes": "이 화면은 일기 저장 화면이 아니다. 다이어리와 분리해 상태 체크 전용 테이블에 저장한다.",
    },
    {
        "id": "screen_mission_06",
        "name": "보호자 미션",
        "screen_file": "MissionView.tsx",
        "uc": "UC-05, UC-06",
        "overview": "보호자가 오늘 생성된 케어 미션을 확인하고 완료 처리한다.",
        "route": "대시보드 > 보호자 미션",
        "active_nav": "홈",
        "image": "06_mission.png",
        "features": [
            {"name": "오늘 미션 조회", "action": "보호자가 미션 화면을 연다.", "output": "오늘 날짜의 GUARDIAN_MISSIONS를 조회한다."},
            {"name": "미션 제목/내용", "action": "보호자가 수행할 내용을 확인한다.", "output": "mission_title, mission_content, mission_reason을 표시한다."},
            {"name": "미션 완료", "action": "보호자가 완료 버튼을 누른다.", "output": "execution_status를 COMPLETED로 변경하고 completed_at을 저장한다."},
            {"name": "개인맞춤 반영", "action": "보호자 설정이 존재한다.", "output": "USER_CARE_PREFERENCES의 미션 유형과 케어 스타일을 반영한다."},
        ],
        "notes": "임산부에게 민감한 원문을 그대로 보여주지 않고 보호자가 수행할 행동 중심으로 표현한다.",
    },
    {
        "id": "screen_diary_list_07",
        "name": "감정 일기 목록",
        "screen_file": "DiaryView.tsx",
        "uc": "UC-07, UC-08, UC-15",
        "overview": "감정 일기와 양쪽 답변이 완료된 스몰토크 기록을 조회한다.",
        "route": "하단 탭 > 다이어리",
        "active_nav": "일기",
        "image": "07_diary_list.png",
        "features": [
            {"name": "일기 목록", "action": "사용자가 저장된 일기 카드를 확인한다.", "output": "DIARY_LOGS를 날짜순으로 표시한다."},
            {"name": "감정 표시", "action": "사용자가 카드의 감정 이모지와 텍스트를 확인한다.", "output": "selected_emotion과 AI 분석 결과를 함께 표현한다."},
            {"name": "상세 접기/펼치기", "action": "사용자가 일기 카드를 누른다.", "output": "일기 본문과 이미지가 펼쳐진다."},
            {"name": "스몰토크 표시", "action": "양쪽 답변이 모두 존재한다.", "output": "SMALL_TALK_ANSWERS를 조합해 다이어리에 표시한다."},
        ],
        "notes": "스몰토크는 양쪽 답변이 모두 있을 때만 다이어리에 노출한다.",
    },
    {
        "id": "screen_diary_write_08",
        "name": "감정 일기 작성",
        "screen_file": "DiaryEntryForm.tsx",
        "uc": "UC-07, UC-08, UC-09, UC-10",
        "overview": "사용자가 감정, 본문, 이미지를 입력하고 AI 감정 분석 결과와 함께 저장한다.",
        "route": "다이어리 > 일기 작성",
        "active_nav": "일기",
        "image": "08_diary_write.png",
        "features": [
            {"name": "본문 입력", "action": "사용자가 오늘의 감정 일기를 작성한다.", "output": "diary_content로 저장된다."},
            {"name": "직접 감정 선택", "action": "사용자가 감정 버튼을 선택한다.", "output": "selected_emotion으로 저장하고 화면에 텍스트로 표시한다."},
            {"name": "AI 자동 감정 분석", "action": "사용자가 분석 버튼을 누른다.", "output": "AI_ANALYSIS_RESULTS.detected_emotion에 저장할 감정을 반환한다."},
            {"name": "이미지 첨부", "action": "사용자가 이미지를 선택한다.", "output": "서버 uploads 경로와 DIARY_LOGS.image_path에 저장한다."},
            {"name": "일기 저장", "action": "사용자가 저장 버튼을 누른다.", "output": "DIARY_LOGS와 AI_ANALYSIS_RESULTS에 데이터를 저장한다."},
        ],
        "notes": "AI 추천과 가전 추천은 이 화면에서 저장된 다이어리 데이터를 개인화 근거로 사용한다.",
    },
    {
        "id": "screen_ai_recommend_09",
        "name": "AI 맞춤 추천 메인",
        "screen_file": "AIRecommendView.tsx",
        "uc": "UC-09",
        "overview": "주차별 추천, 스트레칭, 콘텐츠 탭을 묶는 AI 추천 컨테이너 화면이다.",
        "route": "대시보드 > AI 맞춤 추천",
        "active_nav": "홈",
        "image": "09_ai_recommend.png",
        "features": [
            {"name": "탭 전환", "action": "사용자가 주차별 추천/스트레칭/콘텐츠 탭을 누른다.", "output": "AIWeeklyRecommendView, AIStretchView, AIContentView를 전환한다."},
            {"name": "서버 추천 조회", "action": "화면 진입 시 사용자 ID 또는 이메일로 조회한다.", "output": "GET /api/ai/weekly-recommendations/{identifier} 결과를 상태에 저장한다."},
            {"name": "임신 주차 동기화", "action": "서버가 pregnancy_week를 반환한다.", "output": "프론트 계산값보다 서버 주차를 우선 사용한다."},
            {"name": "추천 데이터 전달", "action": "조회 결과가 성공한다.", "output": "guide와 contents를 하위 추천 컴포넌트로 전달한다."},
        ],
        "notes": "AI 추천은 단일 화면이 아니라 세부 탭 3개로 분리되어 있으므로 화면설계서에서도 별도 화면으로 관리한다.",
    },
    {
        "id": "screen_ai_weekly_10",
        "name": "주차별 AI 추천",
        "screen_file": "AIWeeklyRecommendView.tsx",
        "uc": "UC-09",
        "overview": "현재 주차와 최근 다이어리 맥락을 반영한 추천 식품, 활동, 주의사항을 표시한다.",
        "route": "AI 맞춤 추천 > 주차별 추천",
        "active_nav": "홈",
        "image": "10_ai_weekly.png",
        "features": [
            {"name": "주차 요약", "action": "사용자가 현재 주차 정보를 확인한다.", "output": "태아 발달 크기, 체중, 핵심 안내를 보여준다."},
            {"name": "추천 식품", "action": "사용자가 식품 추천을 확인한다.", "output": "FOOD 타입 추천을 표시한다."},
            {"name": "권장 활동", "action": "사용자가 활동 추천을 확인한다.", "output": "ACTIVITY 타입 추천을 표시한다."},
            {"name": "주의사항", "action": "사용자가 위험 신호를 확인한다.", "output": "WARNING 타입 추천을 표시한다."},
            {"name": "다이어리 반영", "action": "최근 일기 기반 추천이 존재한다.", "output": "personalization과 savedRecommendationCount를 표시한다."},
        ],
        "notes": "WEEKLY_AI_RECOMMENDATIONS.user_id와 diary_id가 개인화 추천 근거다.",
    },
    {
        "id": "screen_ai_stretch_11",
        "name": "주차별 스트레칭",
        "screen_file": "AIStretchView.tsx",
        "uc": "UC-09",
        "overview": "임신 주차와 상황별로 볼 수 있는 스트레칭 영상 추천 화면이다.",
        "route": "AI 맞춤 추천 > 스트레칭",
        "active_nav": "홈",
        "image": "11_ai_stretch.png",
        "features": [
            {"name": "주차별 영상", "action": "사용자가 현재 주차에 맞는 영상을 선택한다.", "output": "유튜브 기반 스트레칭 콘텐츠 링크를 제공한다."},
            {"name": "상황별 스트레칭", "action": "사용자가 불편 증상 상황을 선택한다.", "output": "허리, 골반, 수면 등 상황별 영상을 표시한다."},
            {"name": "주의 문구", "action": "사용자가 운동 전 안내를 확인한다.", "output": "통증, 출혈, 어지러움 등 중단 기준을 안내한다."},
        ],
        "notes": "의학적 위험 신호가 있으면 영상보다 의료진 상담 안내가 우선이다.",
    },
    {
        "id": "screen_ai_content_12",
        "name": "AI 추천 콘텐츠",
        "screen_file": "AIContentView.tsx",
        "uc": "UC-09, UC-16",
        "overview": "주차별 콘텐츠, 체크리스트, 위험 신호를 카드 형태로 보여준다.",
        "route": "AI 맞춤 추천 > 콘텐츠",
        "active_nav": "홈",
        "image": "12_ai_content.png",
        "features": [
            {"name": "주차별 콘텐츠", "action": "사용자가 주차별 정보 카드를 확인한다.", "output": "CONTENT_WEEKLY 추천을 표시한다."},
            {"name": "체크리스트", "action": "사용자가 체크리스트 카드를 확인한다.", "output": "CONTENT_CHECKLIST 추천을 표시한다."},
            {"name": "위험 신호", "action": "사용자가 위험 신호 카드를 확인한다.", "output": "CONTENT_WARNING 추천을 표시한다."},
            {"name": "출처 표시", "action": "사용자가 정보 출처를 확인한다.", "output": "공공기관/의료기관 기반 source를 표시한다."},
        ],
        "notes": "콘텐츠는 챗봇이 아니라 정보 카드 영역이며, 출처 표기가 중요하다.",
    },
    {
        "id": "screen_info_13",
        "name": "신뢰 정보",
        "screen_file": "InfoView.tsx",
        "uc": "UC-16",
        "overview": "영양, 운동, 정신건강, 태아발달, 수면 정보를 공신력 있는 출처 기반으로 제공한다.",
        "route": "대시보드 > 신뢰 정보",
        "active_nav": "홈",
        "image": "13_info.png",
        "features": [
            {"name": "카테고리 탭", "action": "사용자가 정보 카테고리를 선택한다.", "output": "영양, 운동, 정신건강, 태아발달, 수면 데이터를 필터링한다."},
            {"name": "검증 정보 카드", "action": "사용자가 정보 카드를 확인한다.", "output": "제목, 요약, 출처, 임신 시기를 표시한다."},
            {"name": "의학정보 챗봇", "action": "사용자가 챗봇 상담 영역을 연다.", "output": "API 키 설정 시 생성형 AI 상담으로 확장 가능하다."},
            {"name": "혜택 화면 이동", "action": "사용자가 현재 받을 수 있는 혜택 보기를 누른다.", "output": "PregnancyBenefitsView로 이동한다."},
        ],
        "notes": "챗봇은 제거하지 않고 신뢰 정보 화면 내부 기능으로 유지한다.",
    },
    {
        "id": "screen_benefits_14",
        "name": "정부 혜택 보기",
        "screen_file": "PregnancyBenefitsView.tsx",
        "uc": "UC-16, UC-19",
        "overview": "현재 사용자 임신 상태에 맞는 정부 지원 제도를 카드로 정리한다.",
        "route": "신뢰 정보 > 현재 받을 수 있는 혜택 보기",
        "active_nav": "홈",
        "image": "14_benefits.png",
        "features": [
            {"name": "혜택 카드", "action": "사용자가 받을 수 있는 혜택을 확인한다.", "output": "임신/출산 지원금, 진료비, 아동수당 등 정보를 카드로 보여준다."},
            {"name": "대상 조건", "action": "사용자가 대상 조건을 확인한다.", "output": "지원 대상과 신청 시점을 쉽게 읽을 수 있게 표시한다."},
            {"name": "신청 방법", "action": "사용자가 신청 링크를 누른다.", "output": "정부24, 복지로, 보건복지부 등 원문 페이지로 이동한다."},
            {"name": "상황별 갱신", "action": "사용자의 임신 주차/상태 정보가 존재한다.", "output": "현재 상황에 맞는 혜택을 우선 노출한다."},
        ],
        "notes": "정부 제도는 변경 가능성이 있으므로 출처와 확인일을 문서/데이터에 남긴다.",
    },
    {
        "id": "screen_community_15",
        "name": "커뮤니티",
        "screen_file": "CommunityView.tsx",
        "uc": "UC-12, UC-13, UC-14, UC-15",
        "overview": "게시글, 댓글, 좋아요를 제공하는 사용자 커뮤니티 화면이다.",
        "route": "대시보드 > 커뮤니티",
        "active_nav": "홈",
        "image": "15_community.png",
        "features": [
            {"name": "게시글 목록", "action": "사용자가 게시글 목록을 확인한다.", "output": "COMMUNITY_POSTS와 댓글/좋아요 수를 함께 표시한다."},
            {"name": "게시글 작성", "action": "사용자가 제목과 내용을 입력한다.", "output": "COMMUNITY_POSTS에 새 글을 저장한다."},
            {"name": "댓글 작성", "action": "사용자가 댓글을 입력한다.", "output": "COMMUNITY_COMMENTS에 댓글을 저장한다."},
            {"name": "좋아요 토글", "action": "사용자가 좋아요 버튼을 누른다.", "output": "COMMUNITY_POST_LIKES를 생성 또는 삭제하고 카운트를 갱신한다."},
            {"name": "내 활동 이동", "action": "사용자가 내 글/댓글 영역으로 이동한다.", "output": "MyCommunityView에서 본인 활동을 조회한다."},
        ],
        "notes": "좋아요는 계정당 1회만 가능하도록 post_id, user_id 조합 unique 제약을 둔다.",
    },
    {
        "id": "screen_my_community_16",
        "name": "내 커뮤니티",
        "screen_file": "MyCommunityView.tsx",
        "uc": "UC-15",
        "overview": "사용자가 작성한 글과 댓글 단 게시글을 모아 보여준다.",
        "route": "커뮤니티 > 내 활동",
        "active_nav": "홈",
        "image": "16_my_community.png",
        "features": [
            {"name": "내 글 수", "action": "사용자가 본인 글 수를 확인한다.", "output": "사용자 ID 기준 게시글 수를 집계한다."},
            {"name": "내 댓글 수", "action": "사용자가 본인 댓글 수를 확인한다.", "output": "사용자 ID 기준 댓글 수를 집계한다."},
            {"name": "내 게시글", "action": "사용자가 자신이 작성한 글 목록을 확인한다.", "output": "COMMUNITY_POSTS를 user_id 기준으로 조회한다."},
            {"name": "댓글 단 글", "action": "사용자가 댓글을 단 게시글을 확인한다.", "output": "COMMUNITY_COMMENTS 기반으로 관련 게시글을 조회한다."},
        ],
        "notes": "사용자 활동 추적과 커뮤니티 참여도 확인을 위한 화면이다.",
    },
    {
        "id": "screen_appliance_17",
        "name": "가전제어",
        "screen_file": "ApplianceControlView.tsx",
        "uc": "UC-10, UC-11",
        "overview": "AI 추천 가전 설정을 저장하고 Arduino 시연 장치와 동기화한다.",
        "route": "하단 탭 > 가전제어",
        "active_nav": "가전",
        "image": "17_appliance.png",
        "features": [
            {"name": "가전별 상태", "action": "사용자가 가전별 ON/OFF와 목표값을 확인한다.", "output": "APPLIANCE_SETTINGS에 저장된 현재 설정을 표시한다."},
            {"name": "AI 가전 추천", "action": "사용자가 추천 버튼을 누른다.", "output": "최근 DIARY_LOGS와 날씨 기반 목표 온습도를 계산한다."},
            {"name": "설정 저장", "action": "사용자가 설정을 변경 후 저장한다.", "output": "가족 master_id 기준 APPLIANCE_SETTINGS를 갱신한다."},
            {"name": "Arduino 연결", "action": "사용자가 포트 연결을 시도한다.", "output": "pyserial로 USB Serial 연결 상태를 표시한다."},
            {"name": "시연 동기화", "action": "사용자가 sync 버튼을 누른다.", "output": "LCD, LED, 스텝모터, 가습기 모듈에 명령을 전달한다."},
            {"name": "LG 링크", "action": "사용자가 LG 가전 추천 링크를 누른다.", "output": "LG 공식 가전 사이트로 이동한다."},
        ],
        "notes": "Arduino 시연은 로컬호스트와 USB 연결 기준이며 배포 서버에서는 직접 USB 제어가 불가능하다.",
    },
    {
        "id": "screen_smalltalk_18",
        "name": "스몰토크",
        "screen_file": "SmalltalkView.tsx",
        "uc": "UC-15",
        "overview": "임산부와 보호자가 같은 질문에 각각 답변하고 양쪽 답변 완료 시 공개한다.",
        "route": "대시보드 > 스몰토크",
        "active_nav": "일기",
        "image": "18_smalltalk.png",
        "features": [
            {"name": "오늘 질문", "action": "사용자가 오늘의 질문을 확인한다.", "output": "SMALL_TALK_TOPICS에서 질문을 조회한다."},
            {"name": "내 답변 입력", "action": "사용자가 답변을 작성한다.", "output": "SMALL_TALK_ANSWERS에 저장한다."},
            {"name": "상대 답변 상태", "action": "사용자가 상대방 답변 여부를 확인한다.", "output": "상대가 답변 전이면 미완료 상태만 표시한다."},
            {"name": "답변 공개", "action": "두 사용자 답변이 모두 존재한다.", "output": "상대 답변과 내 답변을 함께 표시한다."},
            {"name": "다이어리 연동", "action": "두 답변이 모두 완료된다.", "output": "다이어리 탭에서 해당 날짜 스몰토크 기록으로 표시한다."},
        ],
        "notes": "홈의 스몰토크 기능은 유지하고, 다이어리 저장 표시만 양쪽 답변 완료 조건으로 제한한다.",
    },
    {
        "id": "screen_profile_19",
        "name": "내 정보",
        "screen_file": "ProfileView.tsx",
        "uc": "UC-02, UC-03, UC-06",
        "overview": "사용자 프로필, 연결 코드, 임신 주차, 활동 요약을 확인한다.",
        "route": "하단 탭 > 내 정보",
        "active_nav": "내정보",
        "image": "19_profile.png",
        "features": [
            {"name": "프로필 요약", "action": "사용자가 이름, 역할, 태명 정보를 확인한다.", "output": "USERS 기반 프로필 정보를 표시한다."},
            {"name": "임신 주차", "action": "사용자가 현재 주차를 확인한다.", "output": "pregnancy_start_date 기준 주차를 계산한다."},
            {"name": "연결 코드", "action": "사용자가 연결 코드를 확인한다.", "output": "보호자 연결에 필요한 connection_code를 표시한다."},
            {"name": "활동 요약", "action": "사용자가 커뮤니티/일기 활동을 확인한다.", "output": "사용자별 작성 수와 상태 정보를 표시한다."},
        ],
        "notes": "보호자 계정은 연결된 임산부 정보가 함께 표시되어야 한다.",
    },
    {
        "id": "screen_settings_20",
        "name": "설정",
        "screen_file": "SettingsView.tsx",
        "uc": "UC-03, UC-06",
        "overview": "보호자 케어 선호도와 계정 설정을 관리한다.",
        "route": "하단 탭 > 설정",
        "active_nav": "설정",
        "image": "20_settings.png",
        "features": [
            {"name": "케어 선호 설정", "action": "보호자가 선호 미션 유형을 선택한다.", "output": "USER_CARE_PREFERENCES.preferred_mission_type에 저장한다."},
            {"name": "알림 여부", "action": "보호자가 알림 사용 여부를 변경한다.", "output": "notification_enabled 값을 갱신한다."},
            {"name": "미션 시간", "action": "보호자가 원하는 시간을 입력한다.", "output": "mission_time에 저장한다."},
            {"name": "케어 스타일", "action": "보호자가 말투/방식을 선택한다.", "output": "care_style에 저장하고 미션 생성 시 반영한다."},
            {"name": "로그아웃", "action": "사용자가 로그아웃을 누른다.", "output": "클라이언트 사용자 상태를 초기화한다."},
        ],
        "notes": "케어 설정은 보호자 미션 생성 결과에 직접 영향을 준다.",
    },
    {
        "id": "screen_mental_21",
        "name": "정신건강 케어",
        "screen_file": "MentalCareView.tsx",
        "uc": "UC-16",
        "overview": "임신 중 정서 관리 정보와 안정화 콘텐츠를 제공한다.",
        "route": "대시보드 > 정신건강",
        "active_nav": "홈",
        "image": "21_mental.png",
        "features": [
            {"name": "정서 관리 카드", "action": "사용자가 정신건강 정보를 확인한다.", "output": "공신력 있는 정신건강 안내를 카드로 표시한다."},
            {"name": "자가 점검", "action": "사용자가 감정 상태를 돌아본다.", "output": "불안, 우울, 피로 등 주요 상태를 안내한다."},
            {"name": "상담 안내", "action": "사용자가 위험 신호 문구를 확인한다.", "output": "지속적인 불안/우울은 의료진 또는 상담기관 연결을 권고한다."},
        ],
        "notes": "의학적 진단이 아니라 정보 제공과 상담 권고 화면으로 설계한다.",
    },
    {
        "id": "screen_admin_22",
        "name": "관리자 화면",
        "screen_file": "AdminView.tsx",
        "uc": "UC-20, UC-21, UC-22",
        "overview": "관리자가 회원, 커뮤니티, 접속자, 텍스트 분석, 평균 가전 세팅을 확인한다.",
        "route": "admin 계정 로그인 > 관리자 화면",
        "active_nav": "홈",
        "image": "22_admin.png",
        "features": [
            {"name": "운영 통계", "action": "관리자가 전체 현황을 확인한다.", "output": "회원 수, 게시글 수, 오늘 접속자, 미션 수를 집계한다."},
            {"name": "회원 관리", "action": "관리자가 회원 목록을 확인하거나 삭제한다.", "output": "USERS 및 관련 데이터 삭제 로직을 수행한다."},
            {"name": "대리접속", "action": "관리자가 특정 회원 접속 버튼을 누른다.", "output": "선택한 사용자 화면으로 전환한다."},
            {"name": "커뮤니티 관리", "action": "관리자가 게시글/댓글/좋아요 수를 확인한다.", "output": "COMMUNITY_POSTS, COMMENTS, LIKES를 집계한다."},
            {"name": "커뮤니티 분석", "action": "관리자가 분석 버튼을 누르고 불용어를 입력한다.", "output": "게시글/댓글 텍스트를 분석해 키워드와 워드클라우드 데이터를 표시한다."},
            {"name": "평균 가전 세팅", "action": "관리자가 평균값을 확인한다.", "output": "APPLIANCE_SETTINGS 기반 목표 온도, 습도, 밝기, 풍량 평균을 표시한다."},
        ],
        "notes": "관리자 접근 권한은 admin 계정에 한정한다. Render 무료 인스턴스의 메모리 제한 때문에 AI 모델 상시 적재는 주의한다.",
    },
]


def set_run(run, size=9, bold=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_cell_text(cell, text, size=9, bold=False, color=INK, align=None, after=2):
    p = clear(cell) if not cell.text else cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color)
    return p


def sidebar_item(cell, title, value):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    set_run(run, size=10, bold=True, color=PINK_DARK)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.04
    run = p.add_run(value)
    set_run(run, size=7.4, color=INK)


def add_page(doc, spec, index):
    if index:
        doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    shell = doc.add_table(rows=1, cols=2)
    shell.alignment = WD_TABLE_ALIGNMENT.CENTER
    shell.autofit = False
    set_borders(shell, color=WHITE, size="2")
    left, right = shell.rows[0].cells
    set_width(left, 2650)
    set_width(right, 10500)
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade(left, PINK_SOFT)
    clear(left)
    sidebar_item(left, "화면 ID", spec["id"])
    sidebar_item(left, "화면 이름", spec["name"])
    sidebar_item(left, "관련 유스케이스 ID", spec["uc"])
    sidebar_item(left, "화면 개요", spec["overview"])
    sidebar_item(left, "메뉴 경로", spec["route"])
    sidebar_item(left, "작성자", AUTHOR)

    clear(right)
    title = right.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run(f"{spec['name']} 화면 설계")
    set_run(run, size=18, bold=True, color=PINK)

    shot_table = right.add_table(rows=2, cols=1)
    shot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    shot_table.autofit = False
    set_borders(shot_table, color=GRID, size="5")
    shade(shot_table.rows[0].cells[0], PURPLE_SOFT)
    shade(shot_table.rows[1].cells[0], WHITE)
    add_cell_text(shot_table.rows[0].cells[0], f"화면 설계 - {spec['screen_file']}", size=11.5, bold=True, color=PINK_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    img_path = draw_mock_screen(spec)
    p = clear(shot_table.rows[1].cells[0])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(6.0))

    desc = right.add_table(rows=2, cols=1)
    desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    desc.autofit = False
    set_borders(desc, color=GRID, size="5")
    shade(desc.rows[0].cells[0], PINK)
    shade(desc.rows[1].cells[0], "FFF9FB")
    add_cell_text(desc.rows[0].cells[0], "화면 설명(기능)", size=12, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    body = desc.rows[1].cells[0]
    clear(body)
    feature_table = body.add_table(rows=1, cols=4)
    feature_table.autofit = False
    set_borders(feature_table, color=GRID, size="4")
    headers = ["번호", "기능명", "사용자 액션(Input)", "처리/출력(Output)"]
    widths = [600, 1800, 4050, 4100]
    for i, header in enumerate(headers):
        cell = feature_table.rows[0].cells[i]
        set_width(cell, widths[i])
        shade(cell, PINK_LIGHT)
        add_cell_text(cell, header, size=7.3, bold=True, color=PINK_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    for no, feature in enumerate(spec["features"], start=1):
        row = feature_table.add_row().cells
        values = [no, feature["name"], feature["action"], feature["output"]]
        for i, value in enumerate(values):
            set_width(row[i], widths[i])
            add_cell_text(row[i], value, size=6.8, color=INK, after=0)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("유의사항: ")
    set_run(run, size=7.2, bold=True, color=PINK_DARK)
    run = p.add_run(spec["notes"])
    set_run(run, size=7.0, color=MUTED)


def build():
    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for index, spec in enumerate(SCREENS):
        add_page(doc, spec, index)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
