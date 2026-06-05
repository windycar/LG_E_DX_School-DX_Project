from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "MOMent_화면설계서_프로젝트뷰기준_최종본.docx"
OUTPUT = ROOT / "MOMent_화면설계서_프로젝트뷰기준_최종본_그림삭제_압축수정.docx"

FONT = "맑은 고딕"
PINK_DARK = "A83D5B"
MUTED = "7A6673"


def set_run(run, text, size=7.5, bold=False, color=MUTED):
    run.text = text
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def remove_drawings_from_paragraph(paragraph):
    removed = False
    for drawing in list(paragraph._p.iter(qn("w:drawing"))):
        drawing.getparent().remove(drawing)
        removed = True
    for pict in list(paragraph._p.iter(qn("w:pict"))):
        pict.getparent().remove(pict)
        removed = True
    return removed


def compact_paragraph(paragraph, max_size=7.4):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if run.font.size is None or run.font.size.pt > max_size:
            run.font.size = Pt(max_size)


def compact_document(doc):
    for section in doc.sections:
        section.top_margin = Pt(14)
        section.bottom_margin = Pt(14)
        section.left_margin = Pt(14)
        section.right_margin = Pt(14)

    for paragraph in doc.paragraphs:
        compact_paragraph(paragraph, max_size=8.0)

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            # Do not set fixed row heights. Let Word auto-fit after font reduction.
            for cell in row.cells:
                cell.margin_top = Pt(0)
                cell.margin_bottom = Pt(0)
                cell.margin_left = Pt(1)
                cell.margin_right = Pt(1)
                for paragraph in cell.paragraphs:
                    compact_paragraph(paragraph, max_size=7.0)
                for nested in cell.tables:
                    nested.autofit = True
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            nested_cell.margin_top = Pt(0)
                            nested_cell.margin_bottom = Pt(0)
                            nested_cell.margin_left = Pt(1)
                            nested_cell.margin_right = Pt(1)
                            for paragraph in nested_cell.paragraphs:
                                compact_paragraph(paragraph, max_size=6.6)


def remove_cell_images(cell):
    found = False
    for paragraph in cell.paragraphs:
        if remove_drawings_from_paragraph(paragraph):
            found = True
    for table in cell.tables:
        for row in table.rows:
            for sub_cell in row.cells:
                if remove_cell_images(sub_cell):
                    found = True
    if found and not any(p.text.strip() for p in cell.paragraphs):
        p = cell.paragraphs[0]
        p.alignment = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        set_run(run, "실제 화면 캡처 이미지 삽입 위치", size=6.8, bold=True, color=PINK_DARK)
    return found


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document(SOURCE)

    for paragraph in doc.paragraphs:
        remove_drawings_from_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                remove_cell_images(cell)

    compact_document(doc)

    for rel_id, rel in list(doc.part.rels.items()):
        if "image" in rel.reltype:
            doc.part.drop_rel(rel_id)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
