from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "doc_assets" / "screen_design"
SCREEN_DOC = ROOT / "MOMent_화면설계서.docx"
DB_DOC = ROOT / "MOMent_유스케이스_DB설계서.docx"

FONT = "맑은 고딕"
PINK = "C94E70"
DEEP = "2D1B33"
BLUE = "2E74B5"
LIGHT_PINK = "FDF0F4"
LIGHT_BLUE = "EEF5FF"
LIGHT_GRAY = "F7F8FA"
BORDER = "D9D9D9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10, bold=False, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=18, bold=True, color=PINK)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=14, bold=True, color=DEEP)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=11, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, rows, headers=None, widths=None, header_fill=LIGHT_PINK, font_size=8.5):
    table = doc.add_table(rows=1 if headers else 0, cols=len(headers or rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if headers:
        row = table.rows[0]
        set_repeat_table_header(row)
        for i, header in enumerate(headers):
            cell = row.cells[i]
            set_cell_shading(cell, header_fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                set_cell_width(cell, widths[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            set_run_font(run, size=font_size, bold=True, color=DEEP)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                set_cell_width(cell, widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    doc.add_paragraph()
    return table


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=24, bold=True, color=PINK)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=11, color=DEEP)
    p.paragraph_format.space_after = Pt(14)

    meta = [
        ("프로젝트명", "MOMent 임신 케어 서비스"),
        ("문서 목적", "평가자가 화면 의도, 기능 흐름, 데이터 구조를 별도 문서로 확인할 수 있도록 정리"),
        ("작성 기준일", "2026-06-05"),
        ("작성자", "MOMent 3팀"),
    ]
    add_table(doc, meta, headers=["항목", "내용"], widths=[2200, 6800], header_fill=LIGHT_BLUE, font_size=9)
    return doc


SCREEN_GUIDE_ROWS = [
    ("화면ID", "화면의 고유 ID. 웹 파일명과 맞춰 영어+언더바+숫자 형식으로 관리한다."),
    ("화면 이름", "해당 화면명과 화면의 핵심 특징을 기재한다."),
    ("관련 유스케이스ID", "해당 화면과 직접 연결되는 요구사항 또는 유스케이스 ID를 기재한다."),
    ("화면 개요", "해당 화면에서 사용자가 무엇을 확인하고 처리하는지 설명한다."),
    ("메뉴 경로", "해당 화면에 접근할 수 있는 앱 내 이동 경로를 기재한다."),
    ("작성자", "해당 화면 설계 담당 또는 팀명을 기재한다."),
    ("화면 설계", "실제 모바일 웹 화면 구조를 이미지로 배치한다."),
    ("기능명", "버튼, 입력창, 텍스트, 카드, 이미지 등 화면 내 요소를 번호별로 상세 설명한다."),
    ("유의사항", "OS, 기기, API 실패, 권한, 예외 처리 등 화면 구동 시 주의사항을 기재한다."),
]


SCREENS = [
    {
        "id": "screen_home_01",
        "name": "MOMent 시작 화면",
        "uc": "UC_01, UC_02",
        "overview": "서비스 진입점이다. 임산부/보호자가 로그인 또는 회원가입으로 이동하며 MOMent의 신뢰 정보, AI 추천, 커뮤니티, 가전제어 핵심 가치를 한눈에 확인한다.",
        "path": "앱 실행 > 시작 화면",
        "image": "01_home.png",
        "features": [
            ("1", "앱 로고/서비스명", "MOMent 브랜드와 임신 케어 서비스 정체성을 표시한다."),
            ("2", "핵심 기능 태그", "가전, 보호자 케어, AI 추천, 검증 정보, 커뮤니티 기능을 첫 화면에서 인지시킨다."),
            ("3", "로그인 버튼", "기존 사용자 로그인 화면으로 이동한다."),
            ("4", "회원가입 버튼", "신규 사용자 가입 화면으로 이동한다."),
            ("5", "신뢰 출처 문구", "공공기관·의료기관 기반 정보 제공 방향을 안내한다."),
        ],
        "notes": "모바일 웹 기준 430px 폭에서 표시한다. Vercel 배포 주소 접속 시 보호 설정이 걸려 있으면 공개 도메인을 사용해야 한다.",
    },
    {
        "id": "screen_register_02",
        "name": "회원가입 화면",
        "uc": "UC_01",
        "overview": "임산부, 보호자, 관리자 역할에 맞춰 계정을 생성한다. 임산부 계정은 임신 시작일을 입력해 주차 계산과 추천 기준으로 사용한다.",
        "path": "시작 화면 > 회원가입",
        "image": "02_register.png",
        "features": [
            ("1", "역할 선택", "임산부/보호자 역할을 선택하며 역할별 입력 항목과 연결 로직을 분기한다."),
            ("2", "이메일/비밀번호 입력", "로그인 식별자와 인증 정보를 입력한다."),
            ("3", "이름/태명 입력", "사용자 표시명과 대시보드 태명 정보를 저장한다."),
            ("4", "임신 시작 날짜", "오늘 기준 280일 전부터 오늘까지 선택하도록 제한해 미래 임신일 또는 비정상 날짜를 방지한다."),
            ("5", "연결 코드", "보호자 계정이 임산부 계정과 가족 단위로 연결될 때 사용한다."),
            ("6", "가입 완료 버튼", "입력값 검증 후 USERS 테이블에 저장하고 로그인 상태로 전환한다."),
        ],
        "notes": "임신 시작일은 추천·혜택·주차별 정보의 기준이므로 필수 안내 문구를 노출한다. 서버 연결 실패 시 중복 요청을 막고 오류 메시지를 표시한다.",
    },
    {
        "id": "screen_dashboard_03",
        "name": "대시보드",
        "uc": "UC_03, UC_04, UC_05, UC_09, UC_16",
        "overview": "로그인 후 기본 홈이다. 임신 주차, 오늘의 팁, 상태 체크, 보호자 미션, 스몰토크, 신뢰 정보, 커뮤니티, 가전제어로 이동한다.",
        "path": "로그인 > 대시보드",
        "image": "03_dashboard.png",
        "features": [
            ("1", "사용자/주차 요약", "임신 시작일로 계산한 현재 주차와 태명을 표시한다."),
            ("2", "오늘의 팁 슬라이드", "주차별 팁과 받을 수 있는 정부 혜택 추천을 순환 표시한다."),
            ("3", "오늘의 상태 체크 카드", "임산부가 증상과 감정을 입력하는 상태 체크 화면으로 이동한다."),
            ("4", "보호자 미션 카드", "상태 체크 결과로 생성된 보호자 미션 확인 화면으로 이동한다."),
            ("5", "주요 기능 그리드", "감정 일기, AI 추천, 신뢰 정보, 커뮤니티 등 하위 기능으로 이동한다."),
            ("6", "하단 탭", "대시보드, 가전제어, 다이어리, 내 정보, 설정으로 빠르게 이동한다."),
        ],
        "notes": "보호자 계정은 임산부 데이터 기준으로 일부 정보를 조회한다. API 실패 시 정적 기본 팁을 보여주되 저장성 기능은 재시도하도록 안내한다.",
    },
    {
        "id": "screen_status_check_04",
        "name": "오늘의 상태 체크 화면",
        "uc": "UC_04, UC_05, UC_06",
        "overview": "임산부가 오늘의 몸 상태와 감정을 선택하면 PREGNANCY_STATUS_CHECKS에 저장되고 연결된 보호자에게 맞춤 미션이 생성된다.",
        "path": "대시보드 > 오늘의 상태 체크",
        "image": "04_status_check.png",
        "features": [
            ("1", "증상 선택 칩", "피로, 입덧, 통증 등 복수 증상을 선택해 오늘의 신체 상태를 기록한다."),
            ("2", "감정 선택 칩", "불안, 화남, 우울, 안정 등 감정을 선택한다."),
            ("3", "상태 저장 버튼", "상태 체크 데이터를 서버에 저장하고 보호자 미션 생성 API를 호출한다."),
            ("4", "생성 미션 미리보기", "보호자에게 어떤 종류의 케어 미션이 생성됐는지 임산부에게는 요약만 보여준다."),
            ("5", "개인맞춤 설정 반영", "보호자 케어 선호도에 따라 미션 유형과 말투를 조정한다."),
        ],
        "notes": "상태 체크는 일기 저장이 아니다. DIARY_LOGS가 아니라 PREGNANCY_STATUS_CHECKS에 저장하고, 보호자 미션은 GUARDIAN_MISSIONS로 연결한다.",
    },
    {
        "id": "screen_diary_05",
        "name": "감정 일기 화면",
        "uc": "UC_07, UC_08, UC_15",
        "overview": "감정과 일기 내용을 저장하고 AI 감정 분석 결과를 함께 기록한다. 양쪽 사용자가 답변한 스몰토크만 접기/펼치기 형태로 다이어리에 표시한다.",
        "path": "하단 탭 > 다이어리",
        "image": "05_diary.png",
        "features": [
            ("1", "일기 작성 버튼", "감정 일기 입력 화면을 열어 텍스트, 감정, 이미지 선택을 받는다."),
            ("2", "감정 선택 영역", "사용자가 직접 누른 감정을 텍스트와 이모지로 함께 표시한다."),
            ("3", "AI 자동 감정 분석", "일기 텍스트를 가중치 기반 분석 또는 모델 분석으로 판정해 AI_ANALYSIS_RESULTS에 저장한다."),
            ("4", "일기 목록", "DIARY_LOGS 데이터를 날짜순으로 조회한다."),
            ("5", "스몰토크 기록", "임산부와 보호자 양쪽 답변이 모두 존재할 때만 다이어리 탭에 노출한다."),
            ("6", "접기/펼치기", "감정 일기와 스몰토크 상세 내용은 사용자가 해당 카드를 눌렀을 때만 펼친다."),
        ],
        "notes": "DIARY_SYMPTOMS는 사용하지 않는다. 증상은 상태 체크 테이블에 저장하고, 일기는 감정·본문·환경값·이미지 중심으로 저장한다.",
    },
    {
        "id": "screen_ai_recommend_06",
        "name": "AI 맞춤 추천 화면",
        "uc": "UC_09",
        "overview": "임신 주차별 기본 추천에 최근 다이어리 감정/본문을 반영해 추천 식품, 권장 활동, 주의사항, 스트레칭 콘텐츠를 제공한다.",
        "path": "대시보드 > AI 맞춤 추천",
        "image": "06_ai.png",
        "features": [
            ("1", "현재 주차 카드", "임신 시작일 기준 현재 주차와 태아 발달 상태를 표시한다."),
            ("2", "추천 식품", "공신력 있는 임신 영양 기준과 사용자의 최근 일기 맥락을 반영해 식품을 추천한다."),
            ("3", "권장 활동", "주차별 안전 활동과 최근 컨디션 기반 활동을 표시한다."),
            ("4", "주의사항", "주차별 위험 신호와 최근 불편감 기반 유의사항을 표시한다."),
            ("5", "최근 다이어리 반영 표시", "추천이 DIARY_LOGS와 AI_ANALYSIS_RESULTS를 기반으로 개인화됐는지 보여준다."),
            ("6", "콘텐츠 탭", "체크리스트, 위험 신호, 주차별 콘텐츠, 스트레칭 영상으로 확장한다."),
        ],
        "notes": "WEEKLY_AI_RECOMMENDATIONS는 단순 정적 주차 데이터가 아니라 user_id, diary_id, personalized_reason을 저장해 개인화 근거를 남긴다.",
    },
    {
        "id": "screen_info_benefits_07",
        "name": "신뢰 정보·정부 혜택 화면",
        "uc": "UC_16",
        "overview": "영양, 운동, 정신건강, 태아발달, 수면, 정부 혜택 정보를 공공기관·의료기관 출처 기준으로 제공한다.",
        "path": "대시보드 > 신뢰 정보 > 현재 받을 수 있는 혜택 보기",
        "image": "07_info.png",
        "features": [
            ("1", "정보 카테고리 탭", "영양, 운동, 정신건강, 태아발달, 수면, 정부 혜택으로 구분한다."),
            ("2", "검증 정보 카드", "출처, 대상 주차, 핵심 내용을 카드 형태로 표시한다."),
            ("3", "정부 혜택 카드", "임신/출산 시기와 사용자 상태에 맞는 지원 제도를 보여준다."),
            ("4", "출처 링크", "정부24, 복지로, 보건복지부, 의료기관 페이지로 이동한다."),
            ("5", "의학정보 챗봇 상담", "API 키 입력 시 생성형 AI 상담으로 확장 가능한 영역을 유지한다."),
        ],
        "notes": "챗봇은 신뢰 정보 영역 아래에 유지한다. 정부 혜택은 제도 변경 가능성이 있으므로 출처와 확인일을 함께 관리해야 한다.",
    },
    {
        "id": "screen_community_08",
        "name": "커뮤니티 화면",
        "uc": "UC_12, UC_13, UC_14",
        "overview": "임산부와 보호자가 게시글, 댓글, 좋아요로 상호작용한다. 게시글 좋아요는 계정당 1회만 허용한다.",
        "path": "대시보드 > 커뮤니티",
        "image": "08_community.png",
        "features": [
            ("1", "게시글 목록", "COMMUNITY_POSTS를 최신순으로 조회한다."),
            ("2", "작성자/임신 시기 표시", "사용자 역할과 임신 기간 정보를 함께 보여준다."),
            ("3", "게시글 작성", "제목, 내용, 임신 시기를 입력해 게시글을 생성한다."),
            ("4", "댓글 보기/작성", "COMMUNITY_COMMENTS를 조회하고 댓글을 등록한다."),
            ("5", "좋아요 버튼", "COMMUNITY_POST_LIKES에 사용자별 좋아요 상태를 저장하거나 취소한다."),
            ("6", "내 글/내 댓글", "사용자별 커뮤니티 활동 이력을 별도 조회한다."),
        ],
        "notes": "좋아요는 UNIQUE(post_id, user_id) 제약으로 중복을 방지한다. 게시글 삭제 시 댓글과 좋아요를 먼저 삭제한다.",
    },
    {
        "id": "screen_appliance_09",
        "name": "가전제어 화면",
        "uc": "UC_10, UC_11",
        "overview": "감정 일기와 환경 데이터를 바탕으로 에어컨, 가습기, 제습기, 공기청정기, 무드등 설정을 추천하고 Arduino 시연 장치와 USB로 연동한다.",
        "path": "하단 탭 > 가전제어",
        "image": "09_appliance.png",
        "features": [
            ("1", "AI 추천 설정", "최근 DIARY_LOGS의 감정·온습도 값을 바탕으로 최적 온도와 습도를 계산한다."),
            ("2", "가전별 제어 카드", "각 가전의 목표값과 실행 상태를 표시한다."),
            ("3", "설정 저장", "APPLIANCE_SETTINGS에 가족 단위 현재 설정을 저장한다."),
            ("4", "Arduino 연결 상태", "USB 시리얼 포트 연결 여부와 전송 상태를 보여준다."),
            ("5", "시연 동기화 버튼", "저장된 설정을 Arduino UNO R3에 전송해 LED, LCD, 스텝모터로 표현한다."),
            ("6", "LG 가전 구매 유도", "LG 공식 가전 사이트로 이동하는 확장 링크를 제공한다."),
        ],
        "notes": "시연은 로컬 서버와 USB 연결 기준이다. 스텝모터/가습기 모듈은 외부 5V 전원을 사용하고 GND를 공통으로 연결한다.",
    },
    {
        "id": "screen_smalltalk_10",
        "name": "스몰토크 화면",
        "uc": "UC_15",
        "overview": "매일 하나의 질문에 임산부와 보호자가 각각 답변한다. 두 답변이 모두 존재할 때만 다이어리 탭에 해당 스몰토크 기록을 표시한다.",
        "path": "대시보드 > 스몰토크",
        "image": "10_smalltalk.png",
        "features": [
            ("1", "오늘의 질문", "SMALL_TALK_TOPICS에서 질문을 조회한다."),
            ("2", "내 답변 입력", "현재 로그인 사용자의 답변을 SMALL_TALK_ANSWERS에 저장한다."),
            ("3", "상대 답변 상태", "연결된 파트너가 답변했는지 상태만 표시한다."),
            ("4", "답변 저장", "topic_id, user_id, connection_code 기준으로 답변을 저장한다."),
            ("5", "다이어리 표시 조건", "임산부와 보호자 답변이 모두 있을 때만 다이어리 기록으로 노출한다."),
        ],
        "notes": "대시보드의 스몰토크 진입 기능은 유지한다. 단, 다이어리에는 미완성 스몰토크를 저장된 기록처럼 보이지 않게 처리한다.",
    },
    {
        "id": "screen_profile_settings_11",
        "name": "내 정보·설정 화면",
        "uc": "UC_02, UC_06",
        "overview": "사용자 정보, 연결 코드, 보호자 케어 선호도, 알림 여부, 로그아웃 기능을 관리한다.",
        "path": "하단 탭 > 내 정보 또는 설정",
        "image": "11_profile.png",
        "features": [
            ("1", "프로필 요약", "이름, 역할, 태명, 임신 시작일, 현재 주차를 표시한다."),
            ("2", "연결 코드", "보호자와 임산부 계정을 연결할 수 있는 코드를 표시한다."),
            ("3", "보호자 케어 설정", "선호 미션 유형, 알림 여부, 미션 시간, 케어 말투를 저장한다."),
            ("4", "로그아웃", "클라이언트 로그인 상태를 초기화하고 시작 화면으로 이동한다."),
            ("5", "역할별 메뉴 표시", "임산부/보호자/관리자 권한에 따라 노출 메뉴를 조정한다."),
        ],
        "notes": "보호자 설정은 USER_CARE_PREFERENCES와 연결된다. 연결 코드가 없는 사용자는 보호자 미션, 공유 일정, 스몰토크 매칭이 제한된다.",
    },
    {
        "id": "screen_admin_12",
        "name": "관리자 화면",
        "uc": "UC_18",
        "overview": "admin 계정만 접근 가능한 운영 화면이다. 회원 관리, 커뮤니티 관리, 오늘의 접속자, 커뮤니티 텍스트 분석, 평균 가전 세팅을 확인한다.",
        "path": "admin 로그인 > 관리자 화면",
        "image": "12_admin.png",
        "features": [
            ("1", "회원 통계", "전체 회원, 임산부, 보호자, 오늘 접속자를 집계한다."),
            ("2", "회원 목록", "사용자 ID, 역할, 주차, 연결 코드 등을 표시한다."),
            ("3", "관리자 대리접속", "관리자가 특정 사용자 화면으로 이동해 시연·검증을 수행한다."),
            ("4", "커뮤니티 관리", "게시글, 댓글, 좋아요 수를 확인하고 삭제할 수 있다."),
            ("5", "커뮤니티 분석", "게시글/댓글 전체 텍스트를 Kiwi 형태소 분석 후 워드클라우드 데이터로 보여준다."),
            ("6", "불용어 입력", "분석에서 제외할 단어를 하나씩 입력하고 재분석한다."),
            ("7", "평균 가전 세팅", "사용자들이 만족한 목표 온도, 습도, 무드등 밝기, 공기청정 풍량을 집계한다."),
        ],
        "notes": "관리자 판별은 email/name/role이 admin 또는 ADMIN인 경우로 처리한다. Render Free 512MB 환경에서는 AI 모델 상시 적재가 메모리 문제를 만들 수 있어 지연 로딩 또는 경량 분석을 사용한다.",
    },
]


USE_CASES = [
    ("UC_01", "회원가입", "임산부/보호자/관리자 계정 생성", "이메일, 비밀번호, 이름, 역할, 임신 시작일, 연결 코드", "USERS", "가입 완료 후 로그인 상태로 전환"),
    ("UC_02", "로그인/사용자 식별", "서버 인증 후 역할별 화면 분기", "이메일, 비밀번호", "USERS", "admin은 관리자 화면, 일반 사용자는 대시보드 진입"),
    ("UC_03", "대시보드 조회", "주차, 팁, 기능 바로가기, 혜택 추천 표시", "로그인 사용자 정보", "USERS, WEEKLY_AI_RECOMMENDATIONS", "사용자 현재 주차에 맞는 요약 제공"),
    ("UC_04", "오늘의 상태 체크", "임산부 증상/감정 체크 저장", "증상 목록, 감정 목록", "PREGNANCY_STATUS_CHECKS", "상태 체크 저장 후 보호자 미션 생성 트리거"),
    ("UC_05", "보호자 미션", "임산부 상태에 맞는 보호자 행동 미션 생성/완료", "상태 체크 ID, 보호자 설정", "GUARDIAN_MISSIONS, USER_CARE_PREFERENCES", "보호자에게만 구체 미션 노출"),
    ("UC_06", "보호자 케어 설정", "미션 유형, 알림, 시간, 말투 설정", "선호 유형, 알림 여부, 시간, 케어 스타일", "USER_CARE_PREFERENCES", "상태 체크 기반 미션 생성 시 반영"),
    ("UC_07", "감정 일기 작성", "감정, 본문, 이미지, 환경값을 일기 로그로 저장", "선택 감정, 일기 본문, 이미지", "DIARY_LOGS", "가전 추천과 AI 추천의 개인화 근거로 사용"),
    ("UC_08", "AI 감정 분석", "일기 본문을 AI/가중치 방식으로 분석", "일기 텍스트", "AI_ANALYSIS_RESULTS", "분석 감정을 일기와 1:1로 저장"),
    ("UC_09", "AI 맞춤 주차별 추천", "주차별 기본 정보에 최근 다이어리 맥락 반영", "사용자 ID 또는 이메일", "WEEKLY_AI_RECOMMENDATIONS, DIARY_LOGS, AI_ANALYSIS_RESULTS", "추천 근거 diary_id와 personalized_reason 저장"),
    ("UC_10", "AI 가전 추천", "감정 일기와 환경값 기반 목표 온습도/가전 설정 추천", "사용자 ID", "DIARY_LOGS, APPLIANCE_SETTINGS", "가족 단위 현재 가전 설정 저장"),
    ("UC_11", "Arduino 가전 시연", "저장된 가전 설정을 USB 시리얼로 Arduino에 전송", "포트, Baudrate, 가전 설정", "APPLIANCE_SETTINGS", "LCD, LED, 스텝모터, 가습기 모듈로 시연"),
    ("UC_12", "커뮤니티 게시글/댓글", "게시글 작성, 댓글 조회·작성·삭제", "제목, 내용, 임신 시기, 댓글", "COMMUNITY_POSTS, COMMUNITY_COMMENTS", "사용자별 활동 이력 조회 가능"),
    ("UC_13", "커뮤니티 좋아요", "게시글 좋아요 토글, 계정당 1회 제한", "post_id, user_id", "COMMUNITY_POST_LIKES", "UNIQUE(post_id, user_id)로 중복 방지"),
    ("UC_14", "커뮤니티 분석", "게시글/댓글 텍스트를 형태소 분석하고 불용어 제외", "불용어 목록", "COMMUNITY_POSTS, COMMUNITY_COMMENTS", "워드클라우드 및 키워드 빈도 제공"),
    ("UC_15", "스몰토크", "커플 질문 답변 저장 및 양쪽 답변 완료 시 다이어리 표시", "topic_id, answer_content", "SMALL_TALK_TOPICS, SMALL_TALK_ANSWERS, DIARY_LOGS", "미완성 답변은 다이어리에 노출하지 않음"),
    ("UC_16", "신뢰 정보/정부 혜택", "공공기관·의료기관 기반 정보와 혜택 제공", "주차, 사용자 상태", "TRUSTED_PREGNANCY_INFO", "출처 링크와 확인일 관리 필요"),
    ("UC_17", "공유 캘린더", "연결 코드 기반 가족 일정 공유", "일정 유형, 제목, 내용, 날짜", "SHARED_CALENDAR_EVENTS", "임산부/보호자가 같은 connection_code로 일정 공유"),
    ("UC_18", "관리자 운영", "회원/커뮤니티/분석/평균 가전 세팅 관리", "admin 계정", "전체 주요 테이블", "admin 계정만 접근 가능"),
]


TABLE_DEFS = [
    ("USERS", "회원 기본 정보와 역할·가족 연결 기준 테이블", [
        ("user_id", "INT/BIGINT", "PK", "사용자 고유 번호"),
        ("email", "VARCHAR(100)", "UNIQUE", "로그인 ID. admin 계정 판별에도 사용"),
        ("password", "VARCHAR(255)", "", "암호화된 비밀번호"),
        ("name", "VARCHAR(50)", "", "사용자 이름"),
        ("role", "VARCHAR(20)", "", "PREGNANT, GUARDIAN, ADMIN"),
        ("pregnancy_start_date", "DATE", "NULL", "임산부 주차 계산 기준일"),
        ("baby_nickname", "VARCHAR(50)", "NULL", "태명"),
        ("connection_code", "VARCHAR(20~100)", "NULL", "보호자·임산부 연결 코드"),
        ("parent_user_id", "BIGINT/INT", "FK성 참조", "보호자 계정이 바라보는 임산부 사용자 ID"),
    ]),
    ("PREGNANCY_STATUS_CHECKS", "오늘의 상태 체크 저장 테이블. 보호자 미션 생성의 직접 기준", [
        ("status_check_id", "BIGINT", "PK", "상태 체크 ID"),
        ("user_id", "BIGINT", "FK USERS", "임산부 사용자 ID"),
        ("symptoms", "TEXT", "NULL", "선택 증상 JSON 문자열"),
        ("emotions", "TEXT", "NULL", "선택 감정 JSON 문자열"),
        ("created_at", "DATETIME", "", "상태 체크 시각"),
    ]),
    ("GUARDIAN_MISSIONS", "보호자에게 전달되는 개인맞춤 미션", [
        ("mission_id", "BIGINT", "PK", "미션 ID"),
        ("analysis_id", "BIGINT", "FK AI_ANALYSIS_RESULTS NULL", "일기 분석 기반 미션일 때 연결"),
        ("status_check_id", "BIGINT", "FK PREGNANCY_STATUS_CHECKS NULL", "상태 체크 기반 미션일 때 연결"),
        ("user_id", "BIGINT", "FK USERS", "미션을 받는 보호자 ID"),
        ("mission_title", "VARCHAR(200)", "", "미션 제목"),
        ("mission_content", "TEXT", "", "보호자에게 보이는 수행 내용"),
        ("mission_reason", "TEXT", "NULL", "미션 생성 근거"),
        ("mission_type", "VARCHAR(50)", "", "emotional_support, housework, physical_care 등"),
        ("execution_status", "VARCHAR(50)", "", "PENDING, COMPLETED"),
        ("created_at", "DATETIME", "", "생성일"),
        ("completed_at", "DATETIME", "NULL", "완료일"),
    ]),
    ("USER_CARE_PREFERENCES", "보호자 미션 개인맞춤 설정", [
        ("preference_id", "BIGINT", "PK", "설정 ID"),
        ("user_id", "BIGINT", "FK USERS UNIQUE", "보호자 사용자 ID"),
        ("preferred_mission_type", "VARCHAR(50)", "", "선호 미션 유형"),
        ("notification_enabled", "BOOLEAN", "", "알림 사용 여부"),
        ("mission_time", "VARCHAR(20)", "NULL", "희망 미션 시간"),
        ("care_style", "VARCHAR(50)", "", "warm, practical, short 등 말투/방식"),
        ("updated_at", "DATETIME", "", "수정일"),
    ]),
    ("DIARY_LOGS", "감정 일기 저장 테이블. AI 분석, 가전 추천, 주차별 추천의 근거", [
        ("diary_id", "BIGINT/INT", "PK", "일기 ID"),
        ("user_id", "BIGINT/INT", "FK USERS", "작성자 ID"),
        ("temperature_ambient", "DECIMAL(4,2)", "NULL", "체감/환경 온도. 현재는 시연용 랜덤값"),
        ("humidity_ambient", "DECIMAL(4,2)", "NULL", "체감/환경 습도. 현재는 시연용 랜덤값"),
        ("weather_ambient", "VARCHAR(50)", "NULL", "날씨 API 결과 또는 기본값"),
        ("selected_emotion", "VARCHAR(50)", "", "사용자가 직접 선택한 감정"),
        ("stress_level", "INT", "NULL", "스트레스 단계"),
        ("diary_content", "TEXT", "", "일기 본문"),
        ("small_talk_topic_id", "BIGINT", "NULL", "스몰토크와 연결될 경우 주제 ID"),
        ("recorded_at", "TIMESTAMP/DATETIME", "", "기록일"),
        ("image_path", "VARCHAR(255)", "NULL", "첨부 이미지 경로"),
    ]),
    ("AI_ANALYSIS_RESULTS", "감정 일기 AI 분석 결과", [
        ("analysis_id", "BIGINT/INT", "PK", "분석 결과 ID"),
        ("diary_id", "BIGINT/INT", "FK DIARY_LOGS UNIQUE", "분석 대상 일기 ID"),
        ("detected_emotion", "VARCHAR(50)", "", "AI가 판정한 감정"),
        ("analyzed_at", "TIMESTAMP/DATETIME", "", "분석 시각"),
    ]),
    ("WEEKLY_AI_RECOMMENDATIONS", "주차별 기본 추천과 사용자 맞춤 추천 저장 테이블", [
        ("recommendation_id", "BIGINT", "PK", "추천 ID"),
        ("user_id", "BIGINT", "FK USERS NULL", "NULL이면 공통 기본 추천, 값이 있으면 개인 추천"),
        ("diary_id", "BIGINT", "FK DIARY_LOGS NULL", "개인 추천 생성 근거가 된 최근 일기"),
        ("pregnancy_week", "INT", "", "추천 적용 임신 주차"),
        ("recommendation_type", "VARCHAR(50)", "", "FOOD, ACTIVITY, WARNING, CONTENT_WEEKLY 등"),
        ("title", "VARCHAR(200)", "", "추천 제목"),
        ("content", "TEXT", "", "추천 본문 또는 JSON 콘텐츠"),
        ("personalized_reason", "TEXT", "NULL", "최근 다이어리 반영 사유"),
        ("created_at", "DATETIME", "", "추천 생성일"),
    ]),
    ("APPLIANCE_SETTINGS", "가족 단위 현재 가전 설정", [
        ("setting_id", "BIGINT", "PK", "가전 설정 ID"),
        ("user_id", "BIGINT", "FK USERS", "가족 대표 사용자 ID"),
        ("appliance_name", "VARCHAR(50)", "", "aircon, humidifier, airPurifier 등"),
        ("control_command", "TEXT", "", "목표 온도·습도·풍량 등 JSON 명령"),
        ("execution_status", "VARCHAR(20)", "", "ON/OFF 또는 실행 상태"),
    ]),
    ("COMMUNITY_POSTS", "커뮤니티 게시글", [
        ("post_id", "BIGINT", "PK", "게시글 ID"),
        ("user_id", "BIGINT", "FK USERS", "작성자 ID"),
        ("pregnancy_period", "VARCHAR(50)", "NULL", "임신 시기 표시"),
        ("title", "VARCHAR(200)", "NULL", "제목"),
        ("content", "TEXT", "NULL", "본문"),
        ("created_at", "TIMESTAMP", "", "작성일"),
    ]),
    ("COMMUNITY_COMMENTS", "커뮤니티 댓글", [
        ("comment_id", "BIGINT/INT", "PK", "댓글 ID"),
        ("post_id", "BIGINT/INT", "FK COMMUNITY_POSTS", "대상 게시글"),
        ("user_id", "BIGINT/INT", "FK USERS", "작성자"),
        ("content", "TEXT", "", "댓글 내용"),
        ("created_at", "TIMESTAMP/DATETIME", "", "작성일"),
    ]),
    ("COMMUNITY_POST_LIKES", "커뮤니티 게시글 좋아요", [
        ("like_id", "BIGINT", "PK", "좋아요 ID"),
        ("post_id", "BIGINT", "FK COMMUNITY_POSTS", "좋아요 대상 게시글"),
        ("user_id", "BIGINT", "FK USERS", "좋아요 누른 사용자"),
        ("created_at", "DATETIME", "", "생성일"),
        ("UNIQUE(post_id,user_id)", "제약", "", "계정당 게시글 1회만 좋아요 허용"),
    ]),
    ("SMALL_TALK_TOPICS", "스몰토크 질문", [
        ("topic_id", "BIGINT", "PK", "질문 ID"),
        ("question_text", "VARCHAR(255)", "", "질문 문구"),
        ("created_at", "TIMESTAMP", "", "생성일"),
    ]),
    ("SMALL_TALK_ANSWERS", "스몰토크 답변", [
        ("answer_id", "BIGINT", "PK", "답변 ID"),
        ("topic_id", "BIGINT", "FK SMALL_TALK_TOPICS", "질문 ID"),
        ("user_id", "BIGINT", "FK USERS", "답변자 ID"),
        ("connection_code", "VARCHAR(100)", "NULL", "가족 연결 코드"),
        ("answer_content", "TEXT", "", "답변 본문"),
        ("match_status", "VARCHAR(50)", "NULL", "매칭 상태"),
        ("created_at", "TIMESTAMP", "", "답변일"),
    ]),
    ("SHARED_CALENDAR_EVENTS", "임산부·보호자 공유 일정", [
        ("event_id", "BIGINT/INT", "PK", "일정 ID"),
        ("connection_code", "VARCHAR(100)", "INDEX", "가족 연결 코드"),
        ("event_type", "VARCHAR(50)", "", "산부인과, 태교 여행 등 유형"),
        ("title", "VARCHAR(200)", "", "일정 제목"),
        ("content", "TEXT", "", "일정 내용"),
        ("event_date", "DATE", "", "일정 날짜"),
        ("created_at", "TIMESTAMP/DATETIME", "", "생성일"),
    ]),
    ("TRUSTED_PREGNANCY_INFO", "신뢰 임신 정보·정부 혜택 콘텐츠", [
        ("info_id", "BIGINT", "PK", "정보 ID"),
        ("title", "VARCHAR(200)", "", "정보 제목"),
        ("content", "TEXT", "", "본문"),
        ("source", "VARCHAR(100)", "NULL", "출처"),
        ("pregnancy_period", "VARCHAR(50)", "NULL", "적용 임신 시기"),
    ]),
]


RELATIONSHIPS = [
    ("USERS", "DIARY_LOGS", "1:N", "사용자 한 명은 여러 감정 일기를 작성한다."),
    ("DIARY_LOGS", "AI_ANALYSIS_RESULTS", "1:1", "일기 하나에 하나의 AI 감정 분석 결과를 연결한다."),
    ("USERS", "PREGNANCY_STATUS_CHECKS", "1:N", "임산부 상태 체크 이력을 저장한다."),
    ("PREGNANCY_STATUS_CHECKS", "GUARDIAN_MISSIONS", "1:N", "상태 체크 결과로 보호자 미션을 생성한다."),
    ("USERS", "USER_CARE_PREFERENCES", "1:1", "보호자별 미션 생성 선호 설정을 저장한다."),
    ("USERS", "WEEKLY_AI_RECOMMENDATIONS", "1:N", "사용자별 개인화 추천을 저장한다. NULL user_id는 공통 기본 추천이다."),
    ("DIARY_LOGS", "WEEKLY_AI_RECOMMENDATIONS", "1:N", "최근 일기 내용을 근거로 맞춤 추천을 생성한다."),
    ("USERS", "APPLIANCE_SETTINGS", "1:N", "가족 대표 사용자 기준 현재 가전 설정을 저장한다."),
    ("USERS", "COMMUNITY_POSTS", "1:N", "사용자가 여러 게시글을 작성한다."),
    ("COMMUNITY_POSTS", "COMMUNITY_COMMENTS", "1:N", "게시글 하나에 여러 댓글이 달린다."),
    ("COMMUNITY_POSTS", "COMMUNITY_POST_LIKES", "1:N", "게시글 좋아요를 별도 테이블로 관리한다."),
    ("USERS", "COMMUNITY_POST_LIKES", "1:N", "사용자는 여러 게시글에 좋아요를 누를 수 있으나 같은 글은 1회만 가능하다."),
    ("SMALL_TALK_TOPICS", "SMALL_TALK_ANSWERS", "1:N", "질문 하나에 임산부·보호자 각각의 답변이 저장된다."),
    ("connection_code", "SHARED_CALENDAR_EVENTS", "1:N", "같은 연결 코드를 가진 가족이 일정을 공유한다."),
]


def build_screen_doc():
    doc = setup_doc("MOMent 화면설계서", "현재 React + Vite 모바일 웹앱 기준 화면별 설계 문서")
    add_heading(doc, "1. 화면설계서 작성 기준", 1)
    add_table(doc, SCREEN_GUIDE_ROWS, headers=["항목", "설명"], widths=[1700, 7300], header_fill=LIGHT_BLUE, font_size=8.8)

    add_heading(doc, "2. 화면 목록", 1)
    screen_rows = [(s["id"], s["name"], s["uc"], s["path"]) for s in SCREENS]
    add_table(doc, screen_rows, headers=["화면ID", "화면 이름", "관련 유스케이스ID", "메뉴 경로"], widths=[2100, 2300, 2200, 2400], header_fill=LIGHT_PINK, font_size=8)

    for index, screen in enumerate(SCREENS, start=1):
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_heading(doc, f"3-{index}. {screen['name']}", 1)
        meta_rows = [
            ("화면ID", screen["id"]),
            ("화면 이름", screen["name"]),
            ("관련 유스케이스ID", screen["uc"]),
            ("화면 개요", screen["overview"]),
            ("메뉴 경로", screen["path"]),
            ("작성자", "MOMent 3팀"),
        ]
        add_table(doc, meta_rows, headers=["항목", "설명"], widths=[1700, 7300], header_fill=LIGHT_BLUE, font_size=8.4)

        add_heading(doc, "화면 설계", 2)
        img_path = ASSET_DIR / screen["image"]
        if img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Cm(8.0))
            p.paragraph_format.space_after = Pt(8)
        else:
            add_para(doc, f"화면 이미지 없음: {screen['image']}", color=PINK)

        add_heading(doc, "기능명", 2)
        add_table(doc, screen["features"], headers=["번호", "기능명", "상세 설명"], widths=[800, 2300, 5900], header_fill=LIGHT_PINK, font_size=8.2)

        add_heading(doc, "유의사항", 2)
        add_para(doc, screen["notes"], size=9)

    doc.save(SCREEN_DOC)


def build_db_doc():
    doc = setup_doc("MOMent 유스케이스·DB 설계서", "현재 FastAPI + MySQL + Arduino 시연 연동 기준 기능 및 데이터 설계 문서")

    add_heading(doc, "1. 시스템 개요", 1)
    add_para(
        doc,
        "MOMent는 임산부와 보호자가 함께 사용하는 반응형 모바일 웹앱이다. 임산부는 상태 체크와 감정 일기를 입력하고, 보호자는 임산부 상태에 맞춘 미션을 받는다. "
        "백엔드는 FastAPI와 MySQL을 사용하며, 가전제어 시연은 로컬 환경에서 Arduino UNO R3와 USB Serial로 연결한다.",
        size=9,
    )

    add_heading(doc, "2. 유스케이스 목록", 1)
    add_table(
        doc,
        USE_CASES,
        headers=["ID", "유스케이스명", "목적", "주요 입력", "관련 테이블", "처리 결과"],
        widths=[900, 1700, 2600, 2000, 2300, 2500],
        header_fill=LIGHT_PINK,
        font_size=7.4,
    )

    add_heading(doc, "3. 핵심 업무 흐름", 1)
    flows = [
        ("상태 체크 기반 보호자 미션", "임산부가 오늘의 증상/감정을 저장하면 PREGNANCY_STATUS_CHECKS가 생성되고, 연결된 보호자의 USER_CARE_PREFERENCES를 반영해 GUARDIAN_MISSIONS가 생성된다."),
        ("감정 일기 기반 AI 추천", "DIARY_LOGS와 AI_ANALYSIS_RESULTS를 최근 순으로 분석해 WEEKLY_AI_RECOMMENDATIONS에 user_id, diary_id, personalized_reason을 포함한 맞춤 추천을 저장한다."),
        ("스몰토크 다이어리 표시", "SMALL_TALK_ANSWERS에 임산부와 보호자 답변이 모두 존재할 때만 다이어리 화면의 스몰토크 기록으로 노출한다."),
        ("가전제어 시연", "APPLIANCE_SETTINGS에 저장된 목표 온도, 습도, 풍량, 무드등 정보를 Arduino Serial 명령으로 변환해 LCD, LED, 스텝모터로 표현한다."),
        ("커뮤니티 분석", "COMMUNITY_POSTS와 COMMUNITY_COMMENTS의 텍스트를 형태소 분석하고 불용어를 제외한 키워드 빈도/워드클라우드 데이터를 관리자에게 제공한다."),
    ]
    add_table(doc, flows, headers=["업무 흐름", "설명"], widths=[2300, 6700], header_fill=LIGHT_BLUE, font_size=8.4)

    add_heading(doc, "4. DB 설계 원칙", 1)
    principles = [
        ("중앙 사용자 테이블", "USERS를 기준으로 임산부, 보호자, 관리자 역할을 구분하고 모든 활동 데이터를 사용자와 연결한다."),
        ("상태 체크와 일기 분리", "오늘의 상태 체크는 미션 생성용 데이터이며 PREGNANCY_STATUS_CHECKS에 저장한다. 감정 일기는 DIARY_LOGS에 저장한다."),
        ("DIARY_SYMPTOMS 제거", "현재 프로젝트에서는 DIARY_SYMPTOMS를 사용하지 않는다. 증상은 PREGNANCY_STATUS_CHECKS.symptoms로 관리한다."),
        ("추천 근거 저장", "WEEKLY_AI_RECOMMENDATIONS는 공통 주차 추천과 사용자 맞춤 추천을 함께 관리하며, 맞춤 추천은 user_id와 diary_id로 근거를 남긴다."),
        ("좋아요 중복 방지", "COMMUNITY_POST_LIKES는 UNIQUE(post_id, user_id)로 계정당 1회 좋아요만 허용한다."),
        ("시연 장치 분리", "Arduino는 데이터베이스 테이블이 아니라 APPLIANCE_SETTINGS 값을 읽어 로컬 USB Serial로 표현하는 외부 시연 장치다."),
    ]
    add_table(doc, principles, headers=["원칙", "설명"], widths=[2300, 6700], header_fill=LIGHT_PINK, font_size=8.4)

    add_heading(doc, "5. 테이블 상세 설계", 1)
    for table_name, desc, columns in TABLE_DEFS:
        add_heading(doc, table_name, 2)
        add_para(doc, desc, size=8.8, color=DEEP)
        add_table(doc, columns, headers=["컬럼명", "자료형", "키/제약", "설명"], widths=[2400, 1700, 1800, 3100], header_fill=LIGHT_GRAY, font_size=7.8)

    add_heading(doc, "6. 테이블 관계 설계", 1)
    add_table(doc, RELATIONSHIPS, headers=["기준", "대상", "관계", "설명"], widths=[2100, 2600, 1000, 3300], header_fill=LIGHT_BLUE, font_size=8)

    add_heading(doc, "7. 최근 수정 반영 SQL", 1)
    add_para(doc, "MySQL Workbench에서 먼저 사용할 스키마를 선택하거나 `USE 데이터베이스명;`을 실행한 뒤 적용한다.", size=8.8, bold=True, color=PINK)
    sql = (
        "USE moment_db;\n\n"
        "ALTER TABLE WEEKLY_AI_RECOMMENDATIONS\n"
        "  ADD COLUMN user_id BIGINT NULL,\n"
        "  ADD COLUMN diary_id BIGINT NULL,\n"
        "  ADD COLUMN personalized_reason TEXT NULL,\n"
        "  ADD COLUMN created_at DATETIME NULL DEFAULT CURRENT_TIMESTAMP;\n\n"
        "ALTER TABLE WEEKLY_AI_RECOMMENDATIONS\n"
        "  ADD CONSTRAINT fk_weekly_ai_recommendations_user\n"
        "  FOREIGN KEY (user_id) REFERENCES USERS(user_id),\n"
        "  ADD CONSTRAINT fk_weekly_ai_recommendations_diary\n"
        "  FOREIGN KEY (diary_id) REFERENCES DIARY_LOGS(diary_id);\n\n"
        "DROP TABLE IF EXISTS DIARY_SYMPTOMS;\n\n"
        "CREATE TABLE IF NOT EXISTS COMMUNITY_POST_LIKES (\n"
        "  like_id BIGINT PRIMARY KEY AUTO_INCREMENT,\n"
        "  post_id BIGINT NOT NULL,\n"
        "  user_id BIGINT NOT NULL,\n"
        "  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,\n"
        "  UNIQUE KEY uq_community_post_like_user (post_id, user_id),\n"
        "  INDEX idx_community_post_likes_post_id (post_id),\n"
        "  INDEX idx_community_post_likes_user_id (user_id)\n"
        ");"
    )
    p = doc.add_paragraph()
    run = p.add_run(sql)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)

    add_heading(doc, "8. ERD 반영 메모", 1)
    notes = [
        ("삭제", "DIARY_SYMPTOMS는 현재 기능에서 사용하지 않으므로 ERD와 DB에서 제거한다."),
        ("추가/연결", "WEEKLY_AI_RECOMMENDATIONS.user_id는 USERS.user_id와 연결한다."),
        ("추가/연결", "WEEKLY_AI_RECOMMENDATIONS.diary_id는 DIARY_LOGS.diary_id와 연결해 맞춤 추천 근거를 남긴다."),
        ("유지", "PREGNANCY_STATUS_CHECKS는 오늘의 상태 체크와 보호자 미션 생성을 위해 유지한다."),
        ("유지", "COMMUNITY_POST_LIKES는 좋아요 기능과 관리자 좋아요 수 조회를 위해 유지한다."),
    ]
    add_table(doc, notes, headers=["구분", "반영 내용"], widths=[1400, 7600], header_fill=LIGHT_PINK, font_size=8.4)

    doc.save(DB_DOC)


def main():
    build_screen_doc()
    build_db_doc()
    print(SCREEN_DOC)
    print(DB_DOC)


if __name__ == "__main__":
    main()
