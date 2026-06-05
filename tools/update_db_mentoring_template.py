from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\jyc68\Downloads\5. 5기_DX_3팀_DB멘토링자료 (1).docx")
OUTPUT = ROOT / "5기_DX_3팀_DB멘토링자료_MOMent_DB설계_최신반영.docx"

FONT = "맑은 고딕"


def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(9)


def clone_row(table):
    tbl = table._tbl
    new_tr = deepcopy(table.rows[-1]._tr)
    tbl.append(new_tr)
    return table.rows[-1]


def resize_table(table, row_count):
    while len(table.rows) < row_count:
        clone_row(table)
    while len(table.rows) > row_count:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def replace_table(table, rows):
    resize_table(table, len(rows))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value)


def update_label_value_table(table, values):
    mapping = {row.cells[0].text.strip(): row.cells[1] for row in table.rows[1:]}
    for label, value in values.items():
        if label in mapping:
            set_cell_text(mapping[label], value)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    # Intro paragraphs: keep the template, update the DB overview sentence if present.
    for p in doc.paragraphs:
        if p.text.strip().startswith("MOMent 데이터베이스는 USERS를 중심으로"):
            p.text = (
                "MOMent 데이터베이스는 USERS를 중심으로 회원/역할, 임산부 상태 체크, 보호자 미션, 감정 일기, "
                "다이어리 기반 주차별 AI 추천, 커뮤니티 좋아요, 스몰토크, 공유 일정, 신뢰 정보, 가전 제어 시연 데이터를 연결한다. "
                "이번 최신 반영본은 DIARY_SYMPTOMS 제거, WEEKLY_AI_RECOMMENDATIONS의 user_id/diary_id 연결, "
                "APPLIANCE_SETTINGS의 현재 구조, COMMUNITY_POST_LIKES 중복 제한을 명시한다."
            )

    # Object definition table.
    replace_table(doc.tables[24], [
        ["객체명(테이블)", "객체 설명", "사용 기능"],
        ["USERS", "임산부, 보호자, 관리자 계정을 저장하는 중심 테이블이다. 로그인, 역할 분기, 임신 주차 계산, 보호자 연결의 기준이 된다.", "회원가입, 로그인, 사용자 정보 조회, 프로필 수정, 보호자 연결, 관리자 회원 관리, 임신 주차 계산"],
        ["USER_CARE_PREFERENCES", "보호자 미션을 개인화하기 위한 선호 설정 테이블이다.", "보호자 케어 설정 조회/수정, 보호자 미션 자동 생성"],
        ["PREGNANCY_STATUS_CHECKS", "임산부가 오늘의 상태 체크에서 입력한 증상과 감정을 저장한다. 일기가 아니라 보호자 미션 생성을 위한 입력 테이블이다.", "오늘의 상태 체크 등록, 보호자 미션 생성, 관리자 오늘 활동 사용자 집계"],
        ["GUARDIAN_MISSIONS", "임산부 상태 체크 또는 감정 분석 결과를 보호자가 실천할 수 있는 행동 미션으로 변환하여 저장한다.", "보호자 미션 조회, 완료 처리, 관리자 미션 수 집계"],
        ["DIARY_LOGS", "감정 일기와 환경 데이터를 저장한다. 사용자의 감정 변화, 평균 온습도, 가전 추천, 주차별 AI 맞춤 추천, 관리자 감정 분포 분석의 핵심 데이터다.", "감정 일기 작성/조회, 이미지 기록, 감정 분석, 가전 추천, 주차별 맞춤 추천, 관리자 감정/환경 분석"],
        ["AI_ANALYSIS_RESULTS", "감정 일기 텍스트 분석 결과를 저장한다.", "감정 자동 분석 결과 저장, 다이어리 기반 추천, 보호자 미션 확장"],
        ["APPLIANCE_SETTINGS", "감정과 환경을 기반으로 추천·저장된 가족 단위 가전 제어값을 저장한다. Arduino 시연 동기화의 기준 데이터다.", "가전 추천, bulk 저장, 가족 단위 조회, Arduino 동기화, 관리자 평균 가전 세팅 분석"],
        ["COMMUNITY_POSTS", "임산부/보호자가 작성한 커뮤니티 게시글을 저장한다.", "게시글 작성, 목록 조회, 삭제, 임신 시기 분포, 커뮤니티 텍스트 분석"],
        ["COMMUNITY_COMMENTS", "커뮤니티 게시글에 작성된 댓글을 저장한다.", "댓글 작성/조회/삭제, 댓글 수 집계, 관리자 댓글 관리, 텍스트 분석"],
        ["COMMUNITY_POST_LIKES", "게시글 좋아요를 사용자별로 저장한다. post_id와 user_id 조합을 unique로 관리해 계정당 1회만 허용한다.", "좋아요 토글, 좋아요 수 집계, 내가 좋아요한 글 표시, 관리자 좋아요 수 확인"],
        ["SMALL_TALK_TOPICS", "임산부와 보호자가 함께 답변할 스몰토크 질문을 저장한다.", "오늘 스몰토크 질문 조회"],
        ["SMALL_TALK_ANSWERS", "스몰토크 질문에 대한 사용자 답변을 저장한다. 양쪽 답변이 모두 있을 때만 다이어리 화면에 표시한다.", "스몰토크 답변 저장, 상대 답변 공개 조건 처리, 관리자 오늘 활동 사용자 집계"],
        ["SHARED_CALENDAR_EVENTS", "연결 코드 기준 임산부/보호자 공유 일정을 저장한다.", "공유 일정 등록/조회/삭제, 최근/다음 검진일 계산"],
        ["WEEKLY_AI_RECOMMENDATIONS", "주차별 기본 추천과 최근 다이어리 기반 사용자 맞춤 추천을 함께 저장한다. 맞춤 추천은 user_id, diary_id, personalized_reason으로 근거를 남긴다.", "주차별 AI 추천, 다이어리 기반 개인화, 체크리스트, 위험 신호, 콘텐츠 카드 제공"],
        ["TRUSTED_PREGNANCY_INFO", "공공기관·의료기관 기반 임신 정보와 정부 혜택 콘텐츠를 저장한다.", "신뢰 정보 화면, 정부 혜택 화면, 임신 시기별 정보 제공"],
    ])

    # Appliance table: remove stale analysis_id and mark user_id as required.
    replace_table(doc.tables[37], [
        ["컬럼명", "타입", "키", "NULL", "설명"],
        ["setting_id", "BIGINT", "PK", "N", "가전 설정 ID"],
        ["user_id", "BIGINT", "FK", "N", "가족 master 사용자 ID. 보호자 계정은 연결된 임산부 기준 master_id를 사용한다."],
        ["appliance_name", "VARCHAR(50)", "UNIQUE 조합", "N", "aircon, dehumidifier, humidifier, airPurifier, moodLight 등 가전 구분"],
        ["control_command", "TEXT", "", "N", "JSON 형태 제어 명령. 예: 온도, 습도, 모드, 풍량, 밝기"],
        ["execution_status", "VARCHAR(20)", "", "N", "ON/OFF 또는 실행 상태. 기본값 OFF"],
    ])
    update_label_value_table(doc.tables[38], {
        "사용 기능": "가전 추천, bulk 저장, 가족 단위 조회, Arduino USB Serial 동기화, 관리자 평균 가전 세팅 분석",
        "관계": "USERS.user_id 기준 가족 master_id에 연결된다. 현재 구조에서는 analysis_id 컬럼을 제거하고 사용자·가전명 조합으로 현재 설정을 관리한다.",
    })

    # Weekly AI recommendation table: add new linked columns.
    replace_table(doc.tables[51], [
        ["컬럼명", "타입", "키", "NULL", "설명"],
        ["recommendation_id", "BIGINT", "PK", "N", "추천 ID"],
        ["user_id", "BIGINT", "FK/INDEX", "Y", "맞춤 추천 대상 사용자 ID. NULL이면 전체 사용자 공통 기본 추천"],
        ["diary_id", "BIGINT", "FK/INDEX", "Y", "맞춤 추천 생성 근거가 된 최근 감정 일기 ID"],
        ["pregnancy_week", "INTEGER", "INDEX", "N", "임신 주차. 1~40"],
        ["recommendation_type", "VARCHAR(50)", "INDEX", "N", "META, FOOD, ACTIVITY, WARNING, CONTENT_WEEKLY, CONTENT_CHECKLIST 등"],
        ["title", "VARCHAR(200)", "", "N", "추천 제목"],
        ["content", "TEXT", "", "N", "추천 상세 내용 또는 JSON 콘텐츠"],
        ["personalized_reason", "TEXT", "", "Y", "최근 다이어리 감정/본문을 반영한 추천 사유"],
        ["created_at", "DATETIME", "", "Y", "추천 생성 시간. 기본값 CURRENT_TIMESTAMP"],
    ])
    update_label_value_table(doc.tables[52], {
        "사용 기능": "주차별 AI 추천, 다이어리 기반 개인화 추천, 체크리스트, 위험 신호, 콘텐츠 카드 제공",
        "관계": "USERS.pregnancy_start_date로 pregnancy_week를 계산한다. 공통 추천은 user_id가 NULL이고, 맞춤 추천은 USERS.user_id 및 DIARY_LOGS.diary_id와 연결된다.",
    })

    # Community like relationship detail.
    update_label_value_table(doc.tables[43], {
        "사용 기능": "좋아요 토글, 좋아요 수 집계, 내가 좋아요한 글 표시, 관리자 커뮤니티 좋아요 수 확인",
        "관계": "COMMUNITY_POSTS.post_id 및 USERS.user_id와 N:1로 연결된다. post_id와 user_id 조합에 unique 제약을 둬 계정당 1회만 허용한다.",
    })

    # Relationship table: add/refresh latest relationships.
    replace_table(doc.tables[55], [
        ["관계", "관계 설명", "서비스 의미"],
        ["USERS 1:N DIARY_LOGS", "한 사용자는 여러 감정 일기를 작성한다.", "감정 변화, 환경 데이터, 일기 기록의 기본 관계"],
        ["DIARY_LOGS 1:1 AI_ANALYSIS_RESULTS", "하나의 일기에는 하나의 감정 분석 결과가 연결된다.", "자동 감정 분석과 사용자가 선택한 감정 비교 가능"],
        ["USERS 1:N PREGNANCY_STATUS_CHECKS", "한 임산부는 여러 상태 체크를 등록한다.", "오늘 상태 기반 보호자 미션 생성의 시작점"],
        ["PREGNANCY_STATUS_CHECKS 1:N GUARDIAN_MISSIONS", "상태 체크 하나가 보호자 미션 생성 근거가 된다.", "민감한 상태 정보를 보호자의 행동 미션으로 변환"],
        ["AI_ANALYSIS_RESULTS 1:N GUARDIAN_MISSIONS", "감정 분석 결과는 보호자 미션 확장 근거로 연결될 수 있다.", "일기 기반 미션 확장 가능"],
        ["USERS 1:1 USER_CARE_PREFERENCES", "보호자별 케어 선호를 하나씩 저장한다.", "미션 유형과 표현 방식을 개인화"],
        ["USERS 1:N WEEKLY_AI_RECOMMENDATIONS", "사용자별 맞춤 주차 추천을 저장한다. user_id가 NULL이면 공통 추천이다.", "개인 맞춤 추천과 기본 추천을 같은 테이블에서 관리"],
        ["DIARY_LOGS 1:N WEEKLY_AI_RECOMMENDATIONS", "최근 다이어리 하나가 여러 추천 항목의 근거가 될 수 있다.", "추천 식품, 활동, 주의사항이 실제 일기 맥락과 연결됨"],
        ["USERS 1:N APPLIANCE_SETTINGS", "가족 master 사용자 기준으로 여러 가전 설정을 저장한다.", "임산부와 보호자가 같은 가족 가전 상태를 공유"],
        ["USERS 1:N COMMUNITY_POSTS", "사용자는 여러 게시글을 작성한다.", "커뮤니티 작성자와 임신 시기 표시 가능"],
        ["COMMUNITY_POSTS 1:N COMMUNITY_COMMENTS", "게시글 하나에 여러 댓글이 연결된다.", "게시글별 댓글 수와 댓글 목록 표시"],
        ["COMMUNITY_POSTS 1:N COMMUNITY_POST_LIKES", "게시글 하나에 여러 좋아요가 연결된다.", "좋아요 수 집계"],
        ["USERS 1:N COMMUNITY_POST_LIKES", "사용자는 여러 게시글에 좋아요를 누를 수 있다.", "UNIQUE(post_id, user_id)로 같은 글 중복 좋아요 방지"],
        ["SMALL_TALK_TOPICS 1:N SMALL_TALK_ANSWERS", "질문 하나에 임산부/보호자 답변이 각각 저장된다.", "양쪽 답변 완료 시 다이어리에서 스몰토크 확인"],
        ["USERS.connection_code 1:N SHARED_CALENDAR_EVENTS", "같은 연결 코드를 가진 가족이 공유 일정을 조회한다.", "검진일, 태교 여행, 가족 일정을 함께 관리"],
        ["TRUSTED_PREGNANCY_INFO 독립 참조", "사용자 데이터와 직접 FK를 맺지 않는 검증 정보 콘텐츠다.", "신뢰 정보와 정부 혜택을 임신 시기별로 제공"],
    ])

    # Append explicit changed files table in the same document style.
    doc.add_paragraph("4-4. 최신 DB 반영 파일 명시")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    for idx, text in enumerate(["파일", "반영 내용", "DB 영향"]):
        set_cell_text(hdr[idx], text)
    changed_files = [
        ["backend/models.py", "CommunityPostLike, PregnancyStatusCheck, GuardianMission, UserCarePreference, WeeklyAiRecommendation 최신 모델 구조 기준", "테이블 컬럼 및 FK 설계 기준"],
        ["backend/routers/diary.py", "최근 DIARY_LOGS/AI_ANALYSIS_RESULTS를 분석해 WEEKLY_AI_RECOMMENDATIONS에 user_id, diary_id, personalized_reason 저장", "주차별 추천과 다이어리 로그 연결"],
        ["backend/routers/community.py", "COMMUNITY_POST_LIKES 생성, 좋아요 토글, UNIQUE(post_id, user_id) 적용", "계정당 1회 좋아요 제한"],
        ["backend/routers/appliances.py", "APPLIANCE_SETTINGS에서 analysis_id 제거, user_id+appliance_name 기준 현재 설정 관리, Arduino sync 제공", "가전 설정 테이블 단순화 및 시연 연동"],
        ["backend/routers/guardian.py", "PREGNANCY_STATUS_CHECKS 저장 후 GUARDIAN_MISSIONS 생성, USER_CARE_PREFERENCES 반영", "상태 체크 기반 보호자 미션 흐름"],
        ["backend/routers/admin.py", "회원/커뮤니티/좋아요/오늘 접속자/평균 가전 세팅/커뮤니티 텍스트 분석 조회", "관리자 분석 화면 데이터 기준"],
        ["Project/src/app/AIWeeklyRecommendView.tsx", "최근 다이어리 반영 여부와 저장 추천 수 표시", "맞춤 추천 데이터가 화면에 노출됨"],
        ["Project/src/app/CommunityView.tsx, AdminView.tsx", "커뮤니티 좋아요 및 관리자 좋아요 수 표시", "COMMUNITY_POST_LIKES 조회/집계 활용"],
    ]
    for row in changed_files:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
