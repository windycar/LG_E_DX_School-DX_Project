import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "MOMent_DX_Performance_Tracker_별첨.docx"
ASSET_DIR = ROOT / "doc_assets" / "dx_tracker"
CHART_PATH = ASSET_DIR / "dx_implementation_chart.png"
AI_METRICS_CHART_PATH = ASSET_DIR / "ai_metrics_chart.png"
AI_DATA_CHART_PATH = ASSET_DIR / "ai_dataset_distribution_chart.png"
EVALUATION_PATH = ROOT / "backend" / "diary_emotion_ai" / "reports" / "evaluation.json"
DATASET_SUMMARY_PATH = ROOT / "backend" / "diary_emotion_ai" / "data" / "processed" / "dataset_summary.json"

PINK = "C94E70"
DARK = "2D1B33"
BLUE = "4D8AF0"
LIGHT_PINK = "FCF0F4"
LIGHT_BLUE = "EEF4FF"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "2E9D62"
WHITE = "FFFFFF"


def font_path(bold=False):
    candidates = [
        Path("C:/Windows/Fonts/malgunbd.ttf" if bold else "C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/NanumGothicBold.ttf" if bold else "C:/Windows/Fonts/NanumGothic.ttf"),
    ]
    return next((path for path in candidates if path.exists()), candidates[0])


def make_chart():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 920
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(str(font_path(True)), 54)
    body_font = ImageFont.truetype(str(font_path()), 31)
    bold_font = ImageFont.truetype(str(font_path(True)), 32)
    small_font = ImageFont.truetype(str(font_path()), 25)

    draw.text((90, 55), "MOMent DX 구현 범위", font=title_font, fill=f"#{DARK}")
    draw.text(
        (90, 125),
        "프로젝트 소스와 ERD에서 확인한 디지털 전환 구현 항목 수",
        font=body_font,
        fill=f"#{MID_GRAY}",
    )

    rows = [
        ("데이터 수집·기록", 5, "상태 체크, 감정 일기, 스몰토크, 커뮤니티, 일정"),
        ("분석·개인화", 4, "감정 분석, 주차 추천, 보호자 미션, 커뮤니티 분석"),
        ("서비스 자동화", 4, "미션 생성, 가전 추천, 공공정보 연결, 알림형 대시보드"),
        ("운영·배포", 4, "관리자 화면, MySQL, Vercel/Render, PWA"),
        ("현장 시연 연계", 3, "Arduino Serial, LCD, 모터·LED 제어"),
    ]
    max_value = 5
    chart_left = 420
    chart_right = 1640
    bar_max = chart_right - chart_left
    top = 225
    row_gap = 125
    colors = [PINK, "E8789A", BLUE, "7B68B5", GREEN]

    for index, (label, value, description) in enumerate(rows):
        y = top + index * row_gap
        draw.text((90, y + 3), label, font=bold_font, fill=f"#{DARK}")
        draw.rounded_rectangle(
            (chart_left, y, chart_right, y + 47),
            radius=18,
            fill="#EEF0F4",
        )
        fill_right = chart_left + int(bar_max * value / max_value)
        draw.rounded_rectangle(
            (chart_left, y, fill_right, y + 47),
            radius=18,
            fill=f"#{colors[index]}",
        )
        draw.text((fill_right + 20, y + 3), f"{value}개", font=bold_font, fill=f"#{DARK}")
        draw.text((chart_left, y + 58), description, font=small_font, fill=f"#{MID_GRAY}")

    draw.text(
        (90, 850),
        "※ 본 그래프는 사용자 만족도 조사 결과가 아니라 구현된 기능 범위를 분류한 현황표입니다.",
        font=small_font,
        fill=f"#{MID_GRAY}",
    )
    image.save(CHART_PATH, quality=95)


def make_ai_metrics_chart():
    evaluation = json.loads(EVALUATION_PATH.read_text(encoding="utf-8"))
    per_label = evaluation["per_label"].values()
    label_count = max(len(evaluation["per_label"]), 1)
    metrics = [
        ("정확도", evaluation["accuracy"] * 100),
        ("정밀도", sum(item["precision"] for item in per_label) / label_count * 100),
        ("재현율", sum(item["recall"] for item in evaluation["per_label"].values()) / label_count * 100),
        ("F1 점수", sum(item["f1"] for item in evaluation["per_label"].values()) / label_count * 100),
    ]

    width, height = 1800, 610
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(str(font_path(True)), 48)
    body_font = ImageFont.truetype(str(font_path()), 28)
    bold_font = ImageFont.truetype(str(font_path(True)), 30)
    small_font = ImageFont.truetype(str(font_path()), 23)

    draw.text((90, 48), "AI 감정 분석 모델 내부 검증 결과", font=title_font, fill=f"#{DARK}")
    draw.text(
        (90, 110),
        f"검증 데이터 {evaluation['total']:,}건 기준",
        font=body_font,
        fill=f"#{MID_GRAY}",
    )

    chart_left = 360
    chart_right = 1640
    bar_max = chart_right - chart_left
    top = 190
    row_gap = 82
    colors = [PINK, "E8789A", BLUE, GREEN]

    for index, (label, value) in enumerate(metrics):
        y = top + index * row_gap
        draw.text((90, y + 3), label, font=bold_font, fill=f"#{DARK}")
        draw.rounded_rectangle((chart_left, y, chart_right, y + 42), radius=16, fill="#EEF0F4")
        fill_right = chart_left + int(bar_max * value / 100)
        draw.rounded_rectangle(
            (chart_left, y, fill_right, y + 42),
            radius=16,
            fill=f"#{colors[index]}",
        )
        draw.text((chart_right + 24, y + 1), f"{value:.1f}%", font=bold_font, fill=f"#{DARK}")

    draw.text(
        (90, 545),
        "※ AI Hub 문장을 일기 문체로 변환하고 합성 데이터를 포함한 내부 검증 결과이며, 실제 사용자 데이터 검증이 추가로 필요합니다.",
        font=small_font,
        fill=f"#{MID_GRAY}",
    )
    image.save(AI_METRICS_CHART_PATH, quality=95)


def make_ai_dataset_chart():
    summary = json.loads(DATASET_SUMMARY_PATH.read_text(encoding="utf-8"))
    labels = ["기쁨", "당황", "분노", "불안", "상처", "슬픔", "중립"]
    train_total = summary["train_total"]
    validation_total = summary["validation_total"]
    train_distribution = summary["train_distribution"]
    validation_distribution = summary["validation_distribution"]

    width, height = 1800, 690
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(str(font_path(True)), 48)
    body_font = ImageFont.truetype(str(font_path()), 28)
    bold_font = ImageFont.truetype(str(font_path(True)), 27)
    small_font = ImageFont.truetype(str(font_path()), 22)

    draw.text((90, 48), "감정 분석 학습·검증 데이터 구성", font=title_font, fill=f"#{DARK}")
    draw.text(
        (90, 110),
        f"학습 {train_total:,}건 · 검증 {validation_total:,}건 · 7개 감정 분류",
        font=body_font,
        fill=f"#{MID_GRAY}",
    )

    chart_left = 340
    chart_right = 1500
    bar_max = chart_right - chart_left
    top = 190
    row_gap = 60
    max_ratio = 18

    for index, label in enumerate(labels):
        y = top + index * row_gap
        train_ratio = train_distribution[label] / train_total * 100
        validation_ratio = validation_distribution[label] / validation_total * 100
        draw.text((90, y + 4), label, font=bold_font, fill=f"#{DARK}")
        draw.rounded_rectangle((chart_left, y, chart_right, y + 18), radius=8, fill="#F1F2F5")
        draw.rounded_rectangle(
            (chart_left, y, chart_left + int(bar_max * train_ratio / max_ratio), y + 18),
            radius=8,
            fill=f"#{PINK}",
        )
        draw.rounded_rectangle((chart_left, y + 25, chart_right, y + 43), radius=8, fill="#F1F2F5")
        draw.rounded_rectangle(
            (chart_left, y + 25, chart_left + int(bar_max * validation_ratio / max_ratio), y + 43),
            radius=8,
            fill=f"#{BLUE}",
        )
        draw.text(
            (1520, y - 3),
            f"{train_ratio:.1f}% / {validation_ratio:.1f}%",
            font=small_font,
            fill=f"#{DARK}",
        )

    draw.rounded_rectangle((90, 625, 120, 646), radius=6, fill=f"#{PINK}")
    draw.text((135, 620), "학습", font=small_font, fill=f"#{DARK}")
    draw.rounded_rectangle((240, 625, 270, 646), radius=6, fill=f"#{BLUE}")
    draw.text((285, 620), "검증", font=small_font, fill=f"#{DARK}")
    draw.text(
        (470, 620),
        "※ 중립 감정은 원본 데이터에 없어 합성 데이터로 보완했습니다.",
        font=small_font,
        fill=f"#{MID_GRAY}",
    )
    image.save(AI_DATA_CHART_PATH, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)


def style_table(table, widths_cm, header=True, font_size=8.5):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_column_widths(table, widths_cm)
    if header:
        set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if header and row_index == 0:
                set_cell_shading(cell, PINK)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "맑은 고딕"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
                    run.font.size = Pt(font_size)
                    if header and row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def add_cell_text(cell, text, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.4)
    set_cell_shading(cell, LIGHT_PINK)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(PINK)
    run.font.size = Pt(11)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    return paragraph


def add_picture_centered(doc, path, width_cm):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))


def set_doc_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for level, size, before, after in ((1, 16, 16, 8), (2, 13, 12, 6), (3, 11, 8, 4)):
        style = styles[f"Heading {level}"]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(PINK if level < 3 else DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "맑은 고딕"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("MOMent DX Performance Tracker | DX 3팀")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def build_document():
    make_chart()
    make_ai_metrics_chart()
    make_ai_dataset_chart()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    add_footer(section)
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("IX-1. DX Performance Tracker")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("MOMent 임신 케어 서비스 디지털 전환 성과 별첨")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(PINK)

    add_callout(
        doc,
        "DX 추진 개요",
        "임산부의 상태·감정·임신 주차·커뮤니티 데이터를 수집하고 보호자 미션, 맞춤 추천, "
        "관리자 분석, Arduino 가전 시연으로 연결하는 데이터 기반 임신 케어 서비스이다.",
    )

    add_heading(doc, "1. AS-IS / TO-BE", 1)
    comparison = doc.add_table(rows=1, cols=3)
    headers = ["구분", "AS-IS", "TO-BE"]
    for index, text in enumerate(headers):
        add_cell_text(comparison.rows[0].cells[index], text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("상태 공유", "보호자에게 직접 설명", "상태 체크 후 보호자 미션 생성"),
        ("정보 제공", "사용자가 개별 검색", "임신 주차·상황별 맞춤 추천"),
        ("데이터 활용", "기록 후 활용 어려움", "감정·커뮤니티 데이터를 분석에 활용"),
        ("가전 설정", "사용자가 직접 판단", "추천값 저장 및 Arduino 시연"),
        ("운영 관리", "분산된 정보 확인", "관리자 화면에서 통합 조회·분석"),
    ]
    for row_data in rows:
        cells = comparison.add_row().cells
        for index, text in enumerate(row_data):
            add_cell_text(cells[index], text, bold=index == 0, size=8.7)
            if index == 0:
                set_cell_shading(cells[index], LIGHT_PINK)
    style_table(comparison, [3.0, 6.4, 7.0], header=True, font_size=8.7)

    add_heading(doc, "2. DX 데이터 흐름", 1)
    flow = doc.add_table(rows=1, cols=5)
    flow_items = [
        ("1. 입력", "상태·일기·주차·커뮤니티"),
        ("2. 저장", "FastAPI·MySQL"),
        ("3. 분석", "감정·추천·키워드"),
        ("4. 실행", "미션·정보·가전 설정"),
        ("5. 개선", "관리자 지표 확인"),
    ]
    for index, (label, detail) in enumerate(flow_items):
        cell = flow.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_PINK if index % 2 == 0 else LIGHT_BLUE)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(PINK if index % 2 == 0 else BLUE)
        p2 = cell.add_paragraph(detail)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
    style_table(flow, [3.28] * 5, header=False, font_size=8.3)

    doc.add_page_break()
    add_heading(doc, "3. DX Performance Tracker", 1)
    tracker = doc.add_table(rows=1, cols=4)
    tracker_headers = ["DX 영역", "적용 내용", "확인 지표", "현재 상태"]
    for index, text in enumerate(tracker_headers):
        add_cell_text(tracker.rows[0].cells[index], text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, size=7.7)
    tracker_rows = [
        ("데이터화", "상태·일기·커뮤니티·가전 설정 DB 저장", "핵심 데이터 저장 여부", "구현 완료"),
        ("자동화", "상태 체크 기반 보호자 미션 생성", "입력 후 미션 생성 여부", "구현 완료"),
        ("개인화", "임신 주차·감정 기반 추천", "사용자별 추천 표시 여부", "구현 완료"),
        ("운영 분석", "커뮤니티·환경·가전 데이터 분석", "관리자 분석 화면 제공", "구현 완료"),
        ("모바일 전환", "Vercel·Render 배포 및 PWA 적용", "스마트폰 접속·실행", "구현 완료"),
        ("IoT 연계", "가전 추천값을 Arduino 장치로 전달", "LCD·모터·LED 반응", "시연 구현"),
    ]
    for row_index, row_data in enumerate(tracker_rows):
        cells = tracker.add_row().cells
        for index, text in enumerate(row_data):
            add_cell_text(cells[index], text, bold=index == 0, size=7.5)
            if index == 0:
                set_cell_shading(cells[index], LIGHT_PINK)
            elif row_index % 2 == 1:
                set_cell_shading(cells[index], "FAFAFB")
    style_table(tracker, [2.6, 6.0, 5.0, 2.8], header=True, font_size=8.2)

    add_heading(doc, "4. DX 구현 범위", 1)
    add_picture_centered(doc, CHART_PATH, 15.8)
    add_caption(doc, "[그림 1] MOMent DX 구현 항목 분류")
    add_callout(
        doc,
        "DX 성과 요약",
        "MOMent는 임산부의 일상 데이터를 디지털로 수집하고 AI 분석, 보호자 미션, 맞춤 추천, 관리자 분석, "
        "Arduino 가전 시연까지 연결한 데이터 기반 임신 케어 DX 서비스이다.",
    )

    doc.add_page_break()
    add_heading(doc, "5. AI 분석 데이터 그래프", 1)
    add_picture_centered(doc, AI_METRICS_CHART_PATH, 15.8)
    add_caption(doc, "[그림 2] AI 감정 분석 모델 내부 검증 결과")
    add_picture_centered(doc, AI_DATA_CHART_PATH, 15.8)
    add_caption(doc, "[그림 3] 감정 분석 학습·검증 데이터 구성")

    doc.core_properties.title = "MOMent DX Performance Tracker"
    doc.core_properties.subject = "DX 스쿨 최종 프로젝트 성과 별첨"
    doc.core_properties.author = "팀원/정용철"
    doc.core_properties.keywords = "MOMent, DX, Performance Tracker, 임신 케어"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
