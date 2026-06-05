from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from create_screen_design_full_doc import AUTHOR, SCREENS


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "MOMent_화면설계서_프로젝트뷰기준_최종산출문서용_그림삭제.docx"

FONT = "맑은 고딕"
INK = "2D1B33"
PINK = "C94E70"
PINK_DARK = "A83D5B"
PINK_LIGHT = "FBE7EE"
PINK_SOFT = "FFF7FA"
GRID = "D9C9D1"
GRAY = "F4F4F4"


def set_run(run, size=9, bold=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, size=8.5, bold=False, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=15, bold=True, color=PINK)
    else:
        set_run(run, size=11, bold=True, color=PINK_DARK)
    return p


def add_meta_table(doc, spec):
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    set_table_widths(table, [1800, 7200])
    rows = [
        ("화면 ID", spec["id"]),
        ("화면 이름", spec["name"]),
        ("관련 유스케이스 ID", spec["uc"]),
        ("화면 개요", spec["overview"]),
        ("메뉴 경로", spec["route"]),
        ("작성자", AUTHOR),
    ]
    for idx, (key, value) in enumerate(rows):
        k, v = table.rows[idx].cells
        k.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        v.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(k, PINK_LIGHT)
        set_cell_text(k, key, size=8.2, bold=True, color=PINK_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(v, value, size=8.2)
    doc.add_paragraph()


def add_screen_placeholder(doc, spec):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    set_table_widths(table, [9000])
    shade(table.rows[0].cells[0], PINK_SOFT)
    set_cell_text(
        table.rows[0].cells[0],
        f"화면 설계 - {spec['screen_file']}",
        size=9,
        bold=True,
        color=PINK_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade(table.rows[1].cells[0], GRAY)
    placeholder = (
        "실제 화면 캡처 이미지 삽입 위치\n"
        "최종산출문서에 붙여넣은 뒤, 이 영역에 해당 React 화면 캡처를 삽입한다."
    )
    set_cell_text(table.rows[1].cells[0], placeholder, size=8.3, color="555555", align=WD_ALIGN_PARAGRAPH.CENTER)
    for p in table.rows[1].cells[0].paragraphs:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
    doc.add_paragraph()


def add_feature_table(doc, spec):
    add_heading(doc, "화면 설명(기능)", 2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    set_table_widths(table, [650, 1800, 3550, 3000])
    headers = ["번호", "기능명", "사용자 액션(Input)", "처리/출력(Output)"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, PINK_LIGHT)
        set_cell_text(cell, header, size=7.6, bold=True, color=PINK_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, feature in enumerate(spec["features"], start=1):
        cells = table.add_row().cells
        values = [idx, feature["name"], feature["action"], feature["output"]]
        for c_idx, value in enumerate(values):
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_text(cells[c_idx], value, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else None)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("유의사항: ")
    set_run(r, size=8, bold=True, color=PINK_DARK)
    r = p.add_run(spec["notes"])
    set_run(r, size=8, color="555555")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MOMent 화면설계서")
    set_run(r, size=20, bold=True, color=PINK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("최종산출문서 삽입용 - 그림 삭제 / 복사 안정화 버전")
    set_run(r, size=9, color="555555")

    for index, spec in enumerate(SCREENS, start=1):
        if index > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_heading(doc, f"2-3-{index:02d}. {spec['name']}", 1)
        add_meta_table(doc, spec)
        add_heading(doc, "화면 설계", 2)
        add_screen_placeholder(doc, spec)
        add_feature_table(doc, spec)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
