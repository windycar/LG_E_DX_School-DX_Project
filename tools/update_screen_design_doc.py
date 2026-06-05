from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "doc_assets" / "screen_design"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/malgunbd.ttf") if bold else Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/NanumGothicBold.ttf") if bold else Path("C:/Windows/Fonts/NanumGothic.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


F_TITLE = font(34, True)
F_H2 = font(25, True)
F_BODY = font(20)
F_SMALL = font(16)
F_XS = font(13)


def rounded(draw: ImageDraw.ImageDraw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, xy, text, fnt, fill):
    x, y = xy
    bbox = draw.textbbox((0, 0), text, font=fnt)
    draw.text((x - (bbox[2] - bbox[0]) / 2, y - (bbox[3] - bbox[1]) / 2), text, font=fnt, fill=fill)


def pill(draw, x, y, text, fill="#FCE7F0", color="#C94E70"):
    bbox = draw.textbbox((0, 0), text, font=F_XS)
    w = bbox[2] - bbox[0] + 28
    rounded(draw, (x, y, x + w, y + 28), 14, fill)
    draw.text((x + 14, y + 6), text, font=F_XS, fill=color)
    return w


def draw_phone_base(title: str, subtitle: str = ""):
    img = Image.new("RGB", (430, 860), "#FFF5F7")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 430, 860), fill="#FFF5F7")
    d.text((26, 24), "9:41", font=F_SMALL, fill="#2D1B33")
    d.text((330, 24), "LTE 100%", font=F_XS, fill="#2D1B33")
    d.text((24, 70), title, font=F_H2, fill="#2D1B33")
    if subtitle:
        d.text((24, 104), subtitle, font=F_SMALL, fill="#8A7A83")
    return img, d


def bottom_nav(d, active: str):
    y = 782
    d.rectangle((0, y, 430, 860), fill="#FFFFFF")
    items = [("홈", "⌂"), ("가전", "⚡"), ("일기", "▣"), ("내정보", "◉"), ("설정", "⚙")]
    x = 42
    for name, icon in items:
        color = "#C94E70" if active == name else "#91848D"
        center_text(d, (x, y + 26), icon, F_BODY, color)
        center_text(d, (x, y + 54), name, F_XS, color)
        x += 86


def card(d, x, y, w, h, title, lines=None, icon=None, fill="#FFFFFF", border="#EADCE3"):
    rounded(d, (x, y, x + w, y + h), 22, fill, border, 2)
    if icon:
        d.text((x + 18, y + 16), icon, font=F_BODY, fill="#C94E70")
        tx = x + 54
    else:
        tx = x + 18
    d.text((tx, y + 16), title, font=F_SMALL, fill="#2D1B33")
    for i, line in enumerate(lines or []):
        d.text((tx, y + 44 + i * 24), line, font=F_XS, fill="#6B5D66")


def make_screens():
    screens = []

    def save(name, title, subtitle, draw_fn):
        img, d = draw_phone_base(title, subtitle)
        draw_fn(img, d)
        path = ASSET_DIR / f"{name}.png"
        img.save(path)
        screens.append(path)

    save("01_home", "MOMent", "임신부터 출산까지 함께하는 케어 서비스", lambda img, d: (
        center_text(d, (215, 230), "🌸", font(78), "#C94E70"),
        center_text(d, (215, 330), "MomDal Care", F_BODY, "#C94E70"),
        card(d, 34, 390, 362, 190, "핵심 기능", ["스마트 가전 · 보호자 미션 · AI 추천", "검증 정보 · 커뮤니티 · 감정 일기"], "✨"),
        rounded(d, (54, 620, 376, 672), 20, "#C94E70"),
        center_text(d, (215, 646), "로그인", F_BODY, "#FFFFFF"),
        rounded(d, (54, 690, 376, 742), 20, "#FFFFFF", "#C94E70", 2),
        center_text(d, (215, 716), "회원가입", F_BODY, "#C94E70")
    ))

    save("02_register", "회원가입", "임산부/보호자 역할에 따라 입력 항목 분기", lambda img, d: (
        pill(d, 28, 130, "임산부"), pill(d, 112, 130, "보호자", "#FFFFFF"),
        card(d, 28, 184, 374, 90, "계정 정보", ["이메일, 비밀번호, 이름 입력"], "👤"),
        card(d, 28, 300, 374, 120, "임산부 정보", ["태명 입력", "임신시작 날짜를 적어주세요", "연결 코드 자동 생성"], "🤰"),
        card(d, 28, 448, 374, 100, "보호자 정보", ["임산부 연결 코드 입력", "보호자 미션 수신 준비"], "🧑‍🤝‍🧑"),
        rounded(d, (42, 690, 388, 746), 22, "#C94E70"),
        center_text(d, (215, 718), "회원가입 완료", F_BODY, "#FFFFFF")
    ))

    save("03_dashboard", "대시보드", "오늘 상태와 주차별 팁을 한 화면에서 확인", lambda img, d: (
        card(d, 24, 130, 382, 120, "오늘의 상태 체크", ["감정·증상 입력 시 보호자 미션 생성"], "💗"),
        card(d, 24, 270, 382, 125, "오늘의 팁", ["현재 주차와 임산부 혜택 정보를 슬라이드 추천"], "💡"),
        card(d, 24, 418, 180, 120, "감정 일기", ["AI 감정 분석", "다이어리 저장"], "😊"),
        card(d, 226, 418, 180, 120, "AI 추천", ["주차+최근 일기", "맞춤 추천"], "🤖"),
        card(d, 24, 560, 180, 120, "신뢰 정보", ["의학 정보", "정부 혜택"], "📋"),
        card(d, 226, 560, 180, 120, "커뮤니티", ["게시글·댓글", "좋아요"], "💬"),
        bottom_nav(d, "홈")
    ))

    save("04_status_check", "오늘의 상태 체크", "상태 체크 후 보호자 미션 자동 생성", lambda img, d: (
        card(d, 24, 126, 382, 100, "오늘 느낀 증상", ["입덧 · 피로감 · 허리통증 · 붓기 선택"], "🩺"),
        card(d, 24, 250, 382, 110, "오늘의 감정", ["불안 · 화남 · 우울 · 안정 등 다중 선택"], "💭"),
        card(d, 24, 385, 382, 120, "보호자 미션 생성", ["사용자 맞춤 설정과 상태를 반영", "남편에게 수행할 미션만 전달"], "🎯"),
        rounded(d, (42, 650, 388, 706), 22, "#C94E70"),
        center_text(d, (215, 678), "상태 체크 저장", F_BODY, "#FFFFFF"),
        bottom_nav(d, "홈")
    ))

    save("05_diary", "감정 일기", "텍스트·감정·AI 분석 결과 저장", lambda img, d: (
        card(d, 24, 126, 382, 150, "오늘의 감정 기록", ["직접 감정 선택 또는 AI 자동 감정 분석", "선택된 감정은 텍스트와 함께 표시"], "📝"),
        rounded(d, (42, 300, 388, 430), 18, "#FFFFFF", "#EADCE3", 2),
        d.text((62, 322), "오늘은 몸이 무겁고 잠을 잘 못 잤다.", font=F_SMALL, fill="#2D1B33"),
        rounded(d, (190, 448, 388, 494), 20, "#8B6FD8"),
        center_text(d, (289, 471), "AI 자동 감정 분석", F_SMALL, "#FFFFFF"),
        card(d, 24, 530, 382, 120, "다이어리 저장", ["DIARY_LOGS와 AI_ANALYSIS_RESULTS에 저장", "최근 일기는 AI 추천 개인화에 사용"], "💾"),
        bottom_nav(d, "일기")
    ))

    save("06_ai", "AI 맞춤 추천", "주차별 추천 + 최근 다이어리 반영", lambda img, d: (
        card(d, 24, 126, 382, 108, "현재 임신 주차", ["178주차처럼 비정상 데이터는 시작일 검증으로 차단"], "📅"),
        card(d, 24, 255, 382, 110, "최근 다이어리 반영", ["최근 DIARY_LOGS 5개와 감정 분석 결과 조회", "피로·입덧·수면·불안 키워드 기반 재정렬"], "🤖"),
        card(d, 24, 390, 382, 90, "추천 식품", ["입덧/피로/붓기 상황에 맞는 항목 우선 노출"], "🍎"),
        card(d, 24, 500, 382, 90, "권장 활동", ["산책, 호흡, 스트레칭 등 안전 활동 추천"], "🏃"),
        card(d, 24, 610, 382, 90, "주의 사항", ["공식 정보 기반 경고 신호 안내"], "⚠️"),
        bottom_nav(d, "홈")
    ))

    save("07_info", "신뢰 정보", "공공기관·의학기관 기반 정보 제공", lambda img, d: (
        pill(d, 24, 126, "영양"), pill(d, 90, 126, "운동"), pill(d, 156, 126, "정신건강"), pill(d, 252, 126, "정부혜택"),
        card(d, 24, 178, 382, 120, "검증 정보 카드", ["보건복지부, 질병관리청, 병원 자료 기반", "출처 링크와 쉬운 설명 제공"], "📚"),
        card(d, 24, 320, 382, 120, "현재 받을 수 있는 혜택보기", ["임신/출산 진료비, 첫만남이용권, 부모급여 등", "신청처와 대상 조건 요약"], "🏛️"),
        card(d, 24, 462, 382, 120, "의학정보 챗봇 상담", ["API 키 입력 시 생성형 AI 연동 가능", "공식 지식 기반 답변 정책"], "💬"),
        bottom_nav(d, "홈")
    ))

    save("08_community", "커뮤니티", "게시글·댓글·좋아요와 관리자 분석 연동", lambda img, d: (
        card(d, 24, 126, 382, 100, "게시글 목록", ["임신 주기 태그, 제목, 내용, 작성자 표시"], "💬"),
        card(d, 24, 250, 382, 110, "좋아요 기능", ["COMMUNITY_POST_LIKES로 계정당 1회 제한", "관리자 화면에서도 좋아요 수 확인"], "❤️"),
        card(d, 24, 385, 382, 120, "커뮤니티 분석", ["전체 글/댓글을 Kiwi 분석 후 워드클라우드 표시", "불용어 직접 추가 가능"], "📊"),
        bottom_nav(d, "홈")
    ))

    save("09_appliance", "가전제어", "AI 추천 설정과 Arduino 시연 연동", lambda img, d: (
        card(d, 24, 126, 382, 100, "가전 상태 카드", ["무드등, 에어컨, 가습기, 제습기, 공기청정기"], "🏠"),
        card(d, 24, 250, 382, 120, "AI 기반 가전 추천", ["다이어리 감정과 환경값 기반 제어 추천", "APPLIANCE_SETTINGS에 저장"], "⚙️"),
        card(d, 24, 395, 382, 120, "Arduino USB 시연", ["로컬 백엔드 pyserial로 UNO에 명령 전송", "스텝모터·LED·LCD·가습기 모듈 제어"], "🔌"),
        card(d, 24, 540, 382, 100, "LG 가전 연계", ["LG 공식 가전 사이트로 이동", "확장 구매 유도"], "🛒"),
        bottom_nav(d, "가전")
    ))

    save("10_smalltalk", "스몰토크", "부부가 모두 답변했을 때 다이어리에 표시", lambda img, d: (
        card(d, 24, 126, 382, 110, "오늘의 스몰토크", ["홈에서 스몰토크 기능은 유지", "그날 카드 클릭 시 질문과 답변 표시"], "💕"),
        card(d, 24, 260, 382, 110, "답변 매칭", ["임산부와 보호자가 모두 답변해야 공개", "SMALL_TALK_ANSWERS에 저장"], "🗣️"),
        card(d, 24, 395, 382, 120, "다이어리 연동", ["둘 다 답변한 스몰토크만 다이어리 탭에 표시", "미응답 상태는 저장 목록에서 제외"], "📔"),
        bottom_nav(d, "일기")
    ))

    save("11_profile", "내 정보 / 설정", "연결 코드와 개인 맞춤 설정 관리", lambda img, d: (
        card(d, 24, 126, 382, 110, "프로필", ["태명, 임신 시작일, 연결 코드 확인"], "👤"),
        card(d, 24, 260, 382, 120, "보호자 연결", ["보호자 회원가입 시 연결 코드 입력", "연결된 보호자 정보 표시"], "🔗"),
        card(d, 24, 405, 382, 120, "개인 맞춤 설정", ["미션 선호 유형, 알림 여부, 케어 스타일 저장", "보호자 미션 생성에 반영"], "🛠️"),
        bottom_nav(d, "내정보")
    ))

    save("12_admin", "관리자 화면", "회원·커뮤니티·데이터 분석 관리", lambda img, d: (
        card(d, 24, 126, 382, 100, "회원 관리", ["임산부/보호자/관리자 계정 확인", "관리자 대리 접속 지원"], "👥"),
        card(d, 24, 250, 382, 110, "커뮤니티 관리", ["게시글, 댓글, 좋아요 수 확인 및 삭제"], "🛡️"),
        card(d, 24, 385, 382, 120, "커뮤니티 텍스트 분석", ["Kiwi 기반 키워드 분석", "워드클라우드와 평균 환경 선호값 표시"], "📈"),
        card(d, 24, 530, 382, 100, "데이터 지표", ["오늘 접속자, 감정 분포, 가전 평균 설정"], "📊")
    ))

    return screens


SCREEN_SPECS = [
    ("SC-001", "MOMent 시작/로그인 진입", "Home", "서비스 첫 진입 화면", "로그인 또는 회원가입 화면으로 이동", ["브랜드 메시지 노출", "핵심 기능 태그 표시", "로그인/회원가입 CTA 제공"]),
    ("SC-002", "회원가입", "Register", "임산부와 보호자 역할별 가입 정보 입력", "인증 후 홈 또는 대시보드 이동", ["임산부는 임신 시작일과 태명 입력", "보호자는 연결 코드 입력", "임신 시작일은 280일 전~오늘까지만 허용"]),
    ("SC-003", "대시보드", "Dashboard", "오늘 상태와 주요 기능 진입 허브", "하단 탭 및 카드형 메뉴", ["오늘의 상태 체크", "주차별 팁/혜택 슬라이드", "AI 추천·신뢰정보·커뮤니티 진입"]),
    ("SC-004", "오늘의 상태 체크", "StatusCheck", "임산부 상태 입력 후 보호자 미션 생성", "대시보드 상태체크 카드", ["증상/감정 다중 선택", "PREGNANCY_STATUS_CHECKS 저장", "GUARDIAN_MISSIONS 자동 생성"]),
    ("SC-005", "감정 일기", "Diary", "일기 작성 및 AI 감정 분석", "하단 다이어리 탭", ["DIARY_LOGS 저장", "AI_ANALYSIS_RESULTS 저장", "최근 일기는 AI 추천 개인화에 사용"]),
    ("SC-006", "AI 맞춤 추천", "AIRecommend", "주차별 공식 정보와 최근 일기 기반 추천", "대시보드 AI 추천 카드", ["임신 주차 계산", "최근 다이어리 5개 반영", "WEEKLY_AI_RECOMMENDATIONS 사용자별 저장"]),
    ("SC-007", "신뢰 정보/정부 혜택", "InfoBenefits", "공식 기관 기반 정보와 임산부 혜택 제공", "대시보드 신뢰정보 카드", ["영양/운동/정신건강/태아발달/수면", "정부 혜택 보기", "출처 기반 안내"]),
    ("SC-008", "커뮤니티", "Community", "임산부 커뮤니티와 좋아요 기능", "대시보드 커뮤니티 카드", ["게시글/댓글 작성", "계정당 좋아요 1회", "관리자 분석 데이터로 활용"]),
    ("SC-009", "가전제어", "Appliance", "스마트 가전 추천과 Arduino 시연 연동", "하단 가전제어 탭", ["가전별 ON/OFF 및 설정", "AI 추천 가전 세팅", "Arduino USB 명령 전송"]),
    ("SC-010", "스몰토크", "SmallTalk", "부부 대화 질문과 답변 매칭", "홈/대시보드 스몰토크 진입", ["질문별 답변 저장", "양쪽 모두 답변 시 공개", "다이어리 탭에 스몰토크 기록 표시"]),
    ("SC-011", "내 정보/설정", "ProfileSettings", "연결 코드와 개인 맞춤 설정 관리", "하단 내정보/설정 탭", ["태명/임신 시작일 확인", "보호자 연결 정보", "미션 선호 설정 저장"]),
    ("SC-012", "관리자 화면", "Admin", "회원·커뮤니티·분석 관리", "admin 계정 로그인", ["회원 목록 관리", "커뮤니티 좋아요 확인", "Kiwi 분석/워드클라우드/환경 평균 조회"]),
]


def set_cell_shading(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, size=9, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run.font.color.rgb = RGBColor.from_string(color)


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(8.5)


def add_meta_table(cell, spec):
    screen_id, name, screen_name, overview, route, ui_items = spec
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "맑은 고딕"

    meta = cell.add_table(rows=5, cols=2)
    style_table(meta)
    rows = [
        ("화면 ID", screen_id),
        ("화면 이름", screen_name),
        ("관련 유스케이스 ID", f"UC-{screen_id.split('-')[-1]}"),
        ("화면 개요", overview),
        ("메뉴 경로", route),
    ]
    for idx, (k, v) in enumerate(rows):
        set_cell_text(meta.rows[idx].cells[0], k, True, 8.5)
        set_cell_text(meta.rows[idx].cells[1], v, False, 8.5)
        set_cell_shading(meta.rows[idx].cells[0], "F3F0F7")

    cell.add_paragraph("")
    ui = cell.add_table(rows=len(ui_items) + 1, cols=2)
    style_table(ui)
    set_cell_text(ui.rows[0].cells[0], "UI", True, 8.5)
    set_cell_text(ui.rows[0].cells[1], "화면 설명", True, 8.5)
    set_cell_shading(ui.rows[0].cells[0], "FCE7F0")
    set_cell_shading(ui.rows[0].cells[1], "FCE7F0")
    for i, item in enumerate(ui_items, 1):
        set_cell_text(ui.rows[i].cells[0], str(i), True, 8.5, "C94E70")
        set_cell_text(ui.rows[i].cells[1], item, False, 8.5)


def add_screen_page(doc: Document, image_path: Path, spec):
    screen_id, name, *_ = spec
    title = doc.add_paragraph()
    title.style = doc.styles["Normal"]
    r = title.add_run(f"2-3-{screen_id.split('-')[-1]}. {name}")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "맑은 고딕"

    layout = doc.add_table(rows=1, cols=2)
    layout.style = "Table Grid"
    left, right = layout.rows[0].cells
    left.width = Cm(8)
    right.width = Cm(15)
    set_cell_shading(left, "FFFFFF")
    set_cell_shading(right, "FFFFFF")
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(2.45))
    add_meta_table(right, spec)
    doc.add_page_break()


def remove_between(body, start_el, end_el):
    children = list(body)
    start_idx = children.index(start_el)
    end_idx = children.index(end_el)
    for el in children[start_idx + 1 : end_idx]:
        body.remove(el)


def paragraph_text(el):
    texts = []
    for node in el.iter():
        if node.tag == qn("w:t") and node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def main():
    source = next(ROOT.glob("*최종산출문서 (3).docx"))
    out = ROOT / "5기_DX_3팀_최종산출문서 (3)_화면설계서수정.docx"
    doc = Document(str(source))
    body = doc._body._element

    start_el = None
    end_el = None
    for el in list(body):
        txt = paragraph_text(el)
        if txt == "2-3. 화면 설계서":
            start_el = el
        elif txt.startswith("4. 데이터 요구사항 분석서") and start_el is not None:
            end_el = el
            break
    if start_el is None or end_el is None:
        raise RuntimeError("화면설계서 섹션 위치를 찾지 못했습니다.")

    remove_between(body, start_el, end_el)
    insert_before = end_el
    before_count = len(list(body))

    screens = make_screens()
    intro = doc.add_paragraph()
    r = intro.add_run("MOMent 화면설계서는 모바일 웹앱 사용 흐름을 기준으로 작성하였다. 각 화면은 실제 구현된 React 화면 구조와 백엔드 연동 테이블을 함께 설명한다.")
    r.font.name = "맑은 고딕"
    r.font.size = Pt(10)
    doc.add_paragraph("")

    for path, spec in zip(screens, SCREEN_SPECS):
        add_screen_page(doc, path, spec)

    children_after_append = list(body)
    appended = children_after_append[before_count:]
    for el in appended:
        body.remove(el)
    for el in appended:
        body.insert(body.index(insert_before), el)

    doc.save(str(out))
    print(out)


if __name__ == "__main__":
    main()
